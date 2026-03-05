# rice_analyzer.py
# RICE – Référentiel Intelligent de Compétences Enseignants
# AI engine: extracts a structured competence tree from UE/module fiches (PDF/DOCX)
# Uses NLP techniques for structured information extraction from ESPRIT fiche format

from __future__ import annotations

import io
import re
import uuid
import unicodedata
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple

import os

from fastapi import APIRouter, File, UploadFile, Form, HTTPException
from fastapi.concurrency import run_in_threadpool
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()  # charge .env (DB_NAME, DB_USER, …)

logger = logging.getLogger("rice_analyzer")

import time as _time
import threading as _threading


# ── Startup env-var validation ──────────────────────────────────────────────
_REQUIRED_ENV_VARS: Dict[str, str] = {
    "DB_NAME": "PostgreSQL database name",
    "DB_USER": "PostgreSQL user",
    "DB_PASS": "PostgreSQL password",
}


def _validate_env() -> None:
    """Warn at import-time when required environment variables are unset."""
    for var, desc in _REQUIRED_ENV_VARS.items():
        val = os.getenv(var)
        if not val:
            logger.warning("Environment variable %s (%s) is not set – using default", var, desc)


_validate_env()


# ── Thread-safe cache wrapper ───────────────────────────────────────────────
class _ThreadSafeCache:
    """A generic dict-like cache protected by a threading.Lock.

    All reads and writes go through the lock so concurrent FastAPI
    thread-pool workers never see a partially-updated dict.
    """

    def __init__(self) -> None:
        self._lock = _threading.Lock()
        self._data: Dict[str, Any] = {}
        self._ts: Dict[str, float] = {}  # per-key timestamps

    # --- read -----------------------------------------------------------------
    def get(self, key: str, *, ttl: float = 0) -> Any:
        """Return cached value if it exists and is fresher than *ttl* seconds."""
        with self._lock:
            if key not in self._data:
                return None
            if ttl > 0 and (_time.time() - self._ts.get(key, 0)) >= ttl:
                return None
            return self._data[key]

    # --- write ----------------------------------------------------------------
    def set(self, key: str, value: Any) -> None:
        with self._lock:
            self._data[key] = value
            self._ts[key] = _time.time()

    # --- invalidate -----------------------------------------------------------
    def pop(self, key: str) -> None:
        with self._lock:
            self._data.pop(key, None)
            self._ts.pop(key, None)

    def clear(self) -> None:
        with self._lock:
            self._data.clear()
            self._ts.clear()

    # --- helpers --------------------------------------------------------------
    def keys(self) -> list:
        with self._lock:
            return list(self._data.keys())

    def __bool__(self) -> bool:
        with self._lock:
            return bool(self._data)


# ── Database connection pool ────────────────────────────────────────────────

_DB_POOL: Any = None
_DB_POOL_LOCK = _threading.Lock()


def _get_db_pool():
    """Lazily create and return the shared ThreadedConnectionPool (thread-safe)."""
    global _DB_POOL
    if _DB_POOL is None:
        with _DB_POOL_LOCK:
            if _DB_POOL is None:
                import psycopg2.pool as _pg_pool
                _DB_POOL = _pg_pool.ThreadedConnectionPool(
                    1, 10,
                    dbname=os.getenv("DB_NAME", "d2f"),
                    user=os.getenv("DB_USER", "d2f"),
                    password=os.getenv("DB_PASS", "d2fpasswd"),
                    host=os.getenv("DB_HOST", "localhost"),
                    port=int(os.getenv("DB_PORT", "7432")),
                )
    return _DB_POOL


def _get_db_connection():
    """Acquire a connection from the shared pool.

    Callers MUST eventually call ``_put_db_connection(conn)`` to return the
    connection to the pool (instead of ``conn.close()``).
    """
    return _get_db_pool().getconn()


def _put_db_connection(conn) -> None:
    """Return *conn* to the pool (or close it on pool error)."""
    try:
        _get_db_pool().putconn(conn)
    except Exception:
        try:
            conn.close()
        except Exception:
            pass


# ── Cached enseignant affectations (TTL = 5 min) ────────────────────────────
_AFFECTATIONS_CACHE = _ThreadSafeCache()
_AFFECTATIONS_CACHE_TTL: float = float(os.getenv("RICE_CACHE_TTL", "300"))  # secondes


def _fetch_enseignant_affectations() -> Dict[str, List[str]]:
    """
    Fetch enseignant → savoir-codes mapping dynamically from PostgreSQL.
    Results are cached for 5 minutes to avoid repeated DB queries.
    Returns a dict like {"E001": ["S2a", "C3b", ...], ...}
    Falls back to cached data or empty dict if the DB is unreachable.
    """
    cached = _AFFECTATIONS_CACHE.get("all", ttl=_AFFECTATIONS_CACHE_TTL)
    if cached is not None:
        return cached

    try:
        conn = _get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT e.id, array_agg(s.code ORDER BY s.code)
            FROM enseignant_competences ec
            JOIN enseignants e ON e.id = ec.enseignant_id
            JOIN savoirs     s ON s.id = ec.savoir_id
            GROUP BY e.id
            ORDER BY e.id
        """)
        result: Dict[str, List[str]] = {}
        for ens_id, codes in cur.fetchall():
            result[str(ens_id)] = codes if codes else []
        cur.close()
        _put_db_connection(conn)
        _AFFECTATIONS_CACHE.set("all", result)
        logger.info("Loaded enseignant_affectations from DB: %d enseignants", len(result))
        return result
    except Exception as exc:
        logger.warning("Cannot fetch enseignant_affectations from DB: %s", exc)
        stale = _AFFECTATIONS_CACHE.get("all")  # return stale cache if available
        return stale if stale is not None else {}

_ENS_INFO_CACHE = _ThreadSafeCache()
_ENS_INFO_TTL: float = 300.0


def _fetch_all_enseignants_info() -> Dict[str, "EnseignantInfo"]:
    """Fetch info (id, nom, prenom, modules) for ALL teachers from DB.

    Also loads savoir names from ``enseignant_competences`` as pseudo-modules
    so that ``_match_enseignants_by_module`` has material to work with.
    """
    cached = _ENS_INFO_CACHE.get("all", ttl=_ENS_INFO_TTL)
    if cached is not None:
        return cached

    try:
        conn = _get_db_connection()
        cur = conn.cursor()
        # Base enseignant info
        cur.execute("SELECT id, nom, prenom FROM enseignants")
        res: Dict[str, Any] = {}
        for row in cur.fetchall():
            eid = str(row[0])
            res[eid] = {
                "id": eid,
                "nom": row[1] or "",
                "prenom": row[2] or "",
                "modules": [],
            }
        # Load savoir names as pseudo-modules (task #3)
        try:
            cur.execute("""
                SELECT ec.enseignant_id, s.nom
                FROM enseignant_competences ec
                JOIN savoirs s ON s.id = ec.savoir_id
            """)
            for eid_raw, snom in cur.fetchall():
                eid = str(eid_raw)
                if eid in res and snom:
                    res[eid]["modules"].append(snom)
        except Exception:
            pass  # table may not exist yet
        cur.close()
        _put_db_connection(conn)

        # Convert raw dicts to EnseignantInfo
        info_map: Dict[str, "EnseignantInfo"] = {}
        for eid, d in res.items():
            info_map[eid] = EnseignantInfo(
                id=d["id"], nom=d["nom"], prenom=d["prenom"],
                modules=list(set(d["modules"])),  # dedupe
            )
        _ENS_INFO_CACHE.set("all", info_map)
        return info_map
    except Exception as e:
        logger.error("Failed to fetch enseignants info: %s", e)
        stale = _ENS_INFO_CACHE.get("all")
        return stale if stale is not None else {}


def _dept_to_numeric_id(departement: str) -> int:
    """Map a department code string to its numeric DB id.

    Falls back to 1 (GC) for unknown codes.
    """
    _MAP = {
        "gc": 1, "genie_civil": 1, "genie-civil": 1,
        "info": 2, "informatique": 2,
        "ge": 3, "genie_electrique": 3, "genie-electrique": 3,
        "meca": 4, "genie_mecanique": 4,
        "telecom": 5, "telecommunications": 5,
    }
    return _MAP.get(departement.lower().strip(), 1)


def _create_enseignant_if_new(nom_complet: str, departement: str = "gc") -> Tuple[str, str]:
    """Auto-create a new enseignant row from a name extracted in a fiche module.

    If the name was not fuzzy-matched against any existing DB teacher, this
    function inserts a new row into ``enseignants`` so the extracted professor
    gets a real DB ID that will survive the import filter.

    Returns (new_id, display_name).
    """
    parts = nom_complet.strip().split()
    nom    = parts[0].upper()                             if parts      else "INCONNU"
    prenom = " ".join(parts[1:]).title()                  if len(parts) > 1 else ""

    slug_raw  = re.sub(r"[^A-Z0-9]", "-", f"{nom}-{prenom}".upper())
    slug_base = re.sub(r"-+", "-", slug_raw).strip("-")[:20]
    new_id    = f"EX-{slug_base}"
    mail      = f"{slug_base.lower()[:30]}@esprit.tn"
    display   = f"{prenom} {nom}".strip() if prenom else nom

    try:
        conn = _get_db_connection()
        cur  = conn.cursor()
        cur.execute("""
            INSERT INTO enseignants
                (id, nom, prenom, mail, type, etat, cup, chefdepartement, up_id, dept_id)
            VALUES (%s, %s, %s, %s, 'P', 'A', 'N', 'N', 1, %s)
            ON CONFLICT (id) DO NOTHING
        """, (new_id, nom, prenom, mail, _dept_to_numeric_id(departement)))
        conn.commit()
        cur.close()
        _put_db_connection(conn)
        # Invalidate cache so next load picks up the new row
        _ENS_INFO_CACHE.clear()
        logger.info(f"  Auto-created enseignant from fiche: {new_id} ({display})")
    except Exception as exc:
        logger.warning(f"  Cannot auto-create enseignant '{nom_complet}': {exc}")

    return new_id, display


# ── optional imports (graceful degradation if libs missing) ──────────────────
try:
    import pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

try:
    from docx import Document as DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from rapidfuzz import fuzz as _rfuzz, process as _rprocess
    _FUZZY_OK = True
except ImportError:
    _FUZZY_OK = False

try:
    from sentence_transformers import SentenceTransformer as _SentenceTransformer
    import numpy as _np
    _SEMANTIC_OK = True
except ImportError:
    _SEMANTIC_OK = False

try:
    import ollama as _ollama
    _LLM_OK = True
except ImportError:
    _LLM_OK = False

# ── LLM configuration (Ollama) ────────────────────────────────────────────
_LLM_MODEL  = os.getenv("OLLAMA_MODEL", "mistral")
_OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://localhost:11434")

# ── Authentication / Authorization ────────────────────────────────────────
from fastapi import Depends
from fastapi.security import OAuth2PasswordBearer

_AUTH_ENABLED = os.getenv("RICE_AUTH_ENABLED", "false").lower() in ("true", "1", "yes")
_AUTH_SECRET  = os.getenv("RICE_AUTH_SECRET", "change-me-in-production")

_oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token", auto_error=False)


def _get_current_user(token: Optional[str] = Depends(_oauth2_scheme)) -> Optional[Dict]:
    """Validate the JWT bearer token when auth is enabled.

    When ``RICE_AUTH_ENABLED`` env var is ``false`` (default), authentication is
    skipped and ``None`` is returned (open access – development mode).
    """
    if not _AUTH_ENABLED:
        return None  # auth disabled – allow all
    if not token:
        raise HTTPException(status_code=401, detail="Authentication required (bearer token missing)")
    try:
        import jwt as _pyjwt
        payload = _pyjwt.decode(token, _AUTH_SECRET, algorithms=["HS256"])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token: missing 'sub'")
        return {"id": user_id, "username": payload.get("name", user_id)}
    except ImportError:
        logger.warning("PyJWT not installed – auth check skipped (install PyJWT to enable)")
        return None
    except Exception as exc:
        raise HTTPException(status_code=401, detail=f"Invalid token: {exc}")


rice_router = APIRouter(prefix="/rice", tags=["RICE"])

# ─────────────────────────────────────────────────────────────────────────────
# Pydantic models
# ─────────────────────────────────────────────────────────────────────────────

class EnseignantInfo(BaseModel):
    id: str
    nom: str
    prenom: str
    modules: List[str] = []   # list of module names the teacher handles

class SavoirProposition(BaseModel):
    tmpId: str
    code: str
    nom: str
    description: Optional[str] = None
    type: str                  # THEORIQUE | PRATIQUE
    niveau: str                # N1_DEBUTANT … N5_EXPERT
    enseignantsSuggeres: List[str] = []   # list of enseignant IDs
    refCodes: List[str] = []             # matched referential codes (e.g. S1a, C2b, INFO-A1)

class SousCompetenceProposition(BaseModel):
    tmpId: str
    code: str
    nom: str
    description: Optional[str] = None
    refCodes: List[str] = []         # aggregated from savoirs
    savoirs: List[SavoirProposition] = []

class CompetenceProposition(BaseModel):
    tmpId: str
    code: str
    nom: str
    description: Optional[str] = None
    ordre: int = 1
    refCodes: List[str] = []         # aggregated from sous-compétences
    refDomaine: Optional[str] = None # best domain match (e.g. GC-TECH-S, INFO-A)
    sousCompetences: List[SousCompetenceProposition] = []

class DomaineProposition(BaseModel):
    tmpId: str
    code: str
    nom: str
    description: Optional[str] = None
    refCodes: List[str] = []         # all referential codes found in this domaine
    refDomaine: Optional[str] = None # best domain match (e.g. GC-TECH-S, INFO-A)
    competences: List[CompetenceProposition] = []

class FicheEnseignantExtrait(BaseModel):
    """Professor name extracted from a fiche module."""
    fichier: str                              # source filename
    nom_complet: str                          # raw name as found in fiche
    role: str = "enseignant"                   # responsable | coordinateur | enseignant | intervenant
    matched_id: Optional[str] = None          # matched enseignant ID (if fuzzy-matched)
    matched_nom: Optional[str] = None         # matched full name for display

class RiceAnalysisResult(BaseModel):
    propositions: List[DomaineProposition]
    stats: Dict[str, Any]
    extractedEnseignants: List[FicheEnseignantExtrait] = []  # professors found in fiches
    foundEnseignants: List[EnseignantInfo] = []  # Added this field

# ─────────────────────────────────────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────────────────────────────────────

def _serialize_pdf_tables(tables: list) -> str:
    """
    Convert pdfplumber table data (list-of-rows-of-cells) to 'label: value' text
    so that existing regex patterns (_RE_RESPONSABLE, etc.) can match them.
    Handles both key-value 2-column tables and multi-column header tables.
    """
    lines: List[str] = []
    for table in tables:
        if not table:
            continue
        # Detect header row: first non-empty row with multiple non-empty cells
        header: List[str] = []
        data_rows: List[List[str]] = []
        for i, row in enumerate(table):
            cells = [str(c).strip() if c else "" for c in row]
            if not any(cells):
                continue
            if i == 0 and sum(bool(c) for c in cells) > 1:
                header = cells
            else:
                data_rows.append(cells)

        if header and data_rows:
            # Table with header: emit as "header_col: cell" pairs
            for row in data_rows:
                for col_idx, cell in enumerate(row):
                    if cell and col_idx < len(header) and header[col_idx]:
                        lines.append(f"{header[col_idx]}: {cell}")
        else:
            # Two-column key-value table (label | value)
            for row in table:
                if not row:
                    continue
                cells = [str(c).strip() if c else "" for c in row]
                non_empty = [c for c in cells if c]
                if len(non_empty) == 2:
                    lines.append(f"{non_empty[0]}: {non_empty[1]}")
                elif len(non_empty) > 2:
                    # Flatten multi-column row with pipe separator
                    lines.append(" | ".join(non_empty[:4]))
    return "\n".join(lines)


def _extract_pdf(data: bytes) -> Tuple[str, List]:
    """Return (full_text, raw_tables) from all pages of a PDF.

    ``raw_tables`` is a flat list of pdfplumber tables (each table is a
    list-of-rows; each row is a list of cell strings).  It is passed down
    to :func:`_extract_metadata` so that structured table-cell scanning
    can be used before falling back to regex patterns.
    """
    if not _PDF_OK:
        raise HTTPException(500, "pdfplumber not installed")
    page_texts: List[str] = []
    all_raw_tables: List = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            # Extract structured table data
            try:
                tables = page.extract_tables() or []
                all_raw_tables.extend(tables)          # keep raw for NER
                table_text = _serialize_pdf_tables(tables)
                if table_text:
                    text = f"{text}\n{table_text}" if text else table_text
            except Exception as tbl_err:
                logger.debug(f"Table extraction failed on page: {tbl_err}")
            page_texts.append(text)
    return "\n".join(page_texts), all_raw_tables


def _extract_docx(data: bytes) -> str:
    if not _DOCX_OK:
        raise HTTPException(500, "python-docx not installed")
    doc = DocxDocument(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_text(filename: str, data: bytes) -> Tuple[str, List]:
    """Return ``(text, raw_tables)`` for the given file.

    ``raw_tables`` is only populated for PDF files where pdfplumber can
    extract structured tables; it is an empty list for DOCX / plain-text.
    """
    name = _secure_filename(filename).lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx") or name.endswith(".doc"):
        return _extract_docx(data), []
    # plain text fallback
    return data.decode("utf-8", errors="ignore"), []


def _secure_filename(filename: str) -> str:
    """Sanitise a user-supplied filename to prevent path-traversal attacks.

    Strips directory components, null bytes, and non-ASCII control characters.
    Returns a safe basename suitable for logging and extension checks.
    """
    # Remove null bytes and control chars
    name = re.sub(r'[\x00-\x1f]', '', filename)
    # Take only the base name (strip any directory components)
    name = os.path.basename(name.replace("..", ""))
    # Remove remaining problematic characters
    name = re.sub(r'[<>:"|?*]', '_', name)
    return name.strip() or "unnamed_file"

# ─────────────────────────────────────────────────────────────────────────────
# NLP – Text normalization & utilities
# ─────────────────────────────────────────────────────────────────────────────

def _normalize(text: str) -> str:
    """Strip accents and lowercase for fuzzy matching."""
    return "".join(
        c for c in unicodedata.normalize("NFD", text.lower())
        if unicodedata.category(c) != "Mn"
    )

def _slug(text: str, max_len: int = 30) -> str:
    """Create a short uppercase code from text."""
    words = re.findall(r"[a-zA-Z\u00C0-\u00FF]{3,}", _normalize(text))[:4]
    raw = "-".join(w[:5].upper() for w in words) if words else "ITEM"
    return raw[:max_len]


def _normalize_ref_code(code: str) -> str:
    """Strip any department prefix from a savoir code.

    Examples::

        'GC-01-S2a' → 'S2a'
        'INFO-A1'   → 'A1'
        'S2a'       → 'S2a'
    """
    return re.split(r"[-–]", code)[-1]


def _codes_match(db_code: str, search_code: str) -> bool:
    """Return True if *search_code* matches *db_code*, tolerating dept prefixes.

    Examples::

        _codes_match('GC-01-S2a', 'S2a')  → True
        _codes_match('S2a', 'GC-01-S2a')  → True
        _codes_match('S2a', 'S2a')        → True
        _codes_match('S3b', 'S2a')        → False
    """
    if db_code == search_code:
        return True
    return _normalize_ref_code(db_code) == _normalize_ref_code(search_code)


# ─────────────────────────────────────────────────────────────────────────────
# NLP – Bloom's Taxonomy detection (6 levels → 5 RICE levels)
# ─────────────────────────────────────────────────────────────────────────────

# Bloom level → RICE niveau mapping
_BLOOM_TO_NIVEAU = {
    1: "N1_DEBUTANT",       # Mémoriser / Reconnaître
    2: "N2_ELEMENTAIRE",    # Comprendre
    3: "N3_INTERMEDIAIRE",  # Appliquer
    4: "N4_AVANCE",         # Analyser
    5: "N4_AVANCE",         # Évaluer
    6: "N5_EXPERT",         # Créer
}

# Verb-based Bloom classification (NLP keyword extraction)
# Includes Génie Civil (GC) domain-specific verbs
_BLOOM_VERBS: List[Tuple[re.Pattern, int]] = [
    # Level 6 – Créer (design, build, synthesize)
    (re.compile(
        r"\b(creer|concevoir|developper|produire|construire|elaborer|"
        r"proposer|innover|composer|planifier|mettre\s+en\s+place|realiser|"
        r"dimensionner|rehabiliter|amenager|piloter|optimiser|"
        r"architecturer|programmer|deployer|integrer\s+un\s+systeme|"
        r"usiner|assembler|fabriquer|prototyper|syntheti[sz]er)\b",
        re.I), 6),
    # Level 5 – Évaluer (judge, validate, audit)
    (re.compile(
        r"\b(evaluer|juger|critiquer|justifier|argumenter|"
        r"defendre|recommander|selectionner|"
        r"diagnostiquer|verifier|controler|auditer|expertiser|valider|"
        r"tester|benchmarker|qualifier|certifier)\b", re.I), 5),
    # Level 4 – Analyser (compare, decompose, model)
    (re.compile(
        r"\b(analyser|comparer|distinguer|examiner|"
        r"differencier|decomposer|organiser|categoriser|"
        r"modeliser|interpreter|superviser|instrumenter|calculer|"
        r"debugger?|profiler|tracer|simuler)\b", re.I), 4),
    # Level 3 – Appliquer (execute, configure, use)
    (re.compile(
        r"\b(appliquer|utiliser|manipuler|implementer|"
        r"executer|resoudre|employer|configurer|installer|"
        r"gerer|integrer|regrouper|"
        r"effectuer|rediger|maitriser|tracer|relever|mesurer|"
        r"programmer?|coder|deployer|monter|brancher|connecter)\b", re.I), 3),
    # Level 2 – Comprendre (explain, describe)
    (re.compile(
        r"\b(comprendre|expliquer|decrire|illustrer|"
        r"interpreter|resumer|classifier|discuter|"
        r"se\s+familiariser|reconnaitre\s+les)\b", re.I), 2),
    # Level 1 – Mémoriser (recall, list, name)
    (re.compile(
        r"\b(reconnaitre|identifier|lister|nommer|"
        r"definir|memoriser|citer|rappeler|introduire)\b", re.I), 1),
]

def _detect_bloom_level(text: str) -> int:
    """Detect Bloom's taxonomy level from verb patterns. Returns 1-6."""
    norm = _normalize(text)
    best = 0
    for pattern, level in _BLOOM_VERBS:
        if pattern.search(norm):
            best = max(best, level)
    return best if best > 0 else 2  # default: Comprendre

def _bloom_to_niveau(bloom: int) -> str:
    """Convert Bloom level (1-6) to RICE niveau string."""
    return _BLOOM_TO_NIVEAU.get(bloom, "N2_ELEMENTAIRE")

# ─────────────────────────────────────────────────────────────────────────────
# NLP – Savoir type detection (THEORIQUE / PRATIQUE)
# ─────────────────────────────────────────────────────────────────────────────

_PRATIQUE_PATTERNS = re.compile(
    r"\b(tp|td|travaux\s+pratiques?|travaux\s+diriges?|projet|labo|atelier|"
    r"manipulation|mise\s+en\s+pratique|implementation|realisation|"
    r"developper|developpement|configurer|installer|creer|coder|"
    r"demarrer\s+un\s+projet|generer|manipuler|implementer|"
    r"app\b|mini[\s\-]projet|validation\s+du\s+projet|"
    # Civil / structural engineering
    r"chantier|terrain|essai|laboratoire|releve|mesure|maquette|"
    r"bim|autocad|revit|robot\s+structural|etabs|sap2000|"
    r"modeli|simulation|dimensionn|concevoir|fondation|soutenement|"
    r"terrassement|coffrage|ferraillage|betonnage|topographi|"
    # Computer science / software engineering
    r"programmer?|debugger?|tester|deployer|jenkins|docker|git|"
    r"base\s+de\s+donn[ée]es|reseau\s+local|socket|api\s+rest|"
    r"machine\s+learning|entrainement\s+mod[eè]le|"
    # Electrical / electronic engineering
    r"montage|circuit|oscilloscope|multimetre|banc\s+de\s+test|"
    r"fpga|arduino|raspberry|microcontroleur|"
    # Mechanical engineering
    r"usinage|usinage|fraisage|tournage|moulage|assemblage|"
    r"cfao|catia|solidworks|ansys|"
    # Telecom
    r"wireshark|routeur|switch|vlsi|hfss|matlab\s+simulink)\b",
    re.IGNORECASE,
)

_THEORIQUE_PATTERNS = re.compile(
    r"\b(cours|theorie|theorique|concept|principe|fondement|notion|"
    r"reconnaitre|comprendre|definir|identifier|memoriser|"
    r"introduction|presentation|fondamentaux|"
    # Civil / structural engineering theory
    r"mecanique\s+des\s+sols|resistance\s+des\s+materiaux|rdm|"
    r"bael|eurocode|norme|reglementation|formule|loi\s+de|"
    r"hypothese|demonstration|calcul\s+theorique|"
    # Computer science theory
    r"algorithmique|complexite|automate|grammaire|langage\s+formel|"
    r"protocole|modele\s+osi|architecture\s+logicielle|pattern|"
    # Electrical / electronic theory
    r"loi\s+(de\s+)?ohm|kirchhoff|theor[eè]me|composant|signal|"
    r"transformee|fourier|filtre|amplificateur|transistor|"
    # Mechanical engineering theory
    r"thermodynamique|cin[eé]matique|dynamique|m[eé]canique\s+des\s+fluides|"
    r"mat[eé]riaux|m[eé]tallurgie|tribologie|"
    # Telecom theory
    r"modulation|codage|antenne|propagation|debit|bande\s+passante)\b",
    re.IGNORECASE,
)

class _UniversalPatterns:
    """Department-aware NLP pattern extensions for PRATIQUE/THEORIQUE classification.

    The base global patterns (_PRATIQUE_PATTERNS / _THEORIQUE_PATTERNS) already cover
    keywords for all ESPRIT departments.  This class adds *incremental* dept-specific
    phrases that are too specialised for the base patterns.
    """

    # Additional practical activity patterns per department
    _EXTRA_PRATIQUE: Dict[str, re.Pattern] = {
        "info": re.compile(
            r"\b(entrainer\s+mod[eè]le|fine[\s\-]tuning|notebook\s+jupyter|"
            r"pipeline\s+ci|deployer\s+image|build\s+docker|"  
            r"commit\s+git|merge\s+request|pull\s+request)\b", re.IGNORECASE,
        ),
        "telecom": re.compile(
            r"\b(simulation\s+ns3|simulation\s+opnet|banc\s+rf|"
            r"analyseur\s+spectre|anritsu|configurer\s+routeur|"
            r"wireshark\s+capture|vlsi\s+implementation)\b", re.IGNORECASE,
        ),
        "ge": re.compile(
            r"\b(montage\s+circuit|banc\s+moteur|pupitre\s+electrique|"
            r"tp\s+fpga|tp\s+automate|maquette\s+electrique)\b", re.IGNORECASE,
        ),
        "meca": re.compile(
            r"\b(atelier\s+usinage|atelier\s+fraisage|banc\s+mecatronique|"
            r"tp\s+catia|tp\s+solidworks|maquette\s+robot)\b", re.IGNORECASE,
        ),
    }

    # Additional theory/concept patterns per department
    _EXTRA_THEORIQUE: Dict[str, re.Pattern] = {
        "info": re.compile(
            r"\b(algorithmique\s+avancee|theorie\s+des\s+graphes|"
            r"automate\s+fini|complexite\s+algorithmique|"
            r"paradigme\s+programmation|theorie\s+des\s+langages)\b", re.IGNORECASE,
        ),
        "telecom": re.compile(
            r"\b(theorie\s+(de\s+l[a'])?information|theorie\s+shannon|"
            r"electromagn[eé]tisme\s+th[eé]orique|propagation\s+th[eé]orie|"
            r"calcul\s+bilan\s+liaison)\b", re.IGNORECASE,
        ),
        "ge": re.compile(
            r"\b(theorie\s+(des\s+)?circuits|theorie\s+(de\s+la\s+)?commande|"
            r"electromagnetisme\s+cours|analyse\s+harmonique)\b", re.IGNORECASE,
        ),
        "meca": re.compile(
            r"\b(theorie\s+mecanique|mecanique\s+th[eé]orique|"
            r"cours\s+thermodynamique|cinématique\s+th[eé]orique)\b", re.IGNORECASE,
        ),
    }

    @classmethod
    def score(cls, text: str, department: str) -> Tuple[int, int]:
        """Return (pratique_score, theorique_score) combining base + dept-specific patterns.

        Args:
            text: raw text to classify (normalization applied internally)
            department: ESPRIT department code (e.g. 'gc', 'info', 'ge', 'meca', 'telecom')
        Returns:
            (prat_score, theo_score): higher score determines the type
        """        
        norm = _normalize(text)
        prat = len(_PRATIQUE_PATTERNS.findall(norm))
        theo = len(_THEORIQUE_PATTERNS.findall(norm))
        dept = department.lower().strip()
        extra_p = cls._EXTRA_PRATIQUE.get(dept)
        extra_t = cls._EXTRA_THEORIQUE.get(dept)
        if extra_p:
            prat += len(extra_p.findall(norm))
        if extra_t:
            theo += len(extra_t.findall(norm))
        return prat, theo


def _detect_type(text: str, departement: str = "gc") -> str:
    """Classify a savoir as PRATIQUE or THEORIQUE using department-aware NLP.

    Uses _UniversalPatterns which combines the base cross-department keyword
    patterns with incremental department-specific extensions.
    """
    prat_score, theo_score = _UniversalPatterns.score(text, departement)
    return "PRATIQUE" if prat_score > theo_score else "THEORIQUE"

# ─────────────────────────────────────────────────────────────────────────────
# LLM Layer – Ollama-powered extraction (augments & replaces regex NLP)
# ─────────────────────────────────────────────────────────────────────────────

_LLM_TIMEOUT = int(os.getenv("RICE_LLM_TIMEOUT", "90"))  # seconds


def _escape_prompt(text: str) -> str:
    """Strip characters that could confuse or inject into LLM prompts."""
    # Remove potential JSON-breaking chars and control sequences
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    # Collapse excessive whitespace
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text


def _llm_chat(messages: List[Dict], temperature: float = 0.05) -> Optional[str]:
    """Send a chat request to the local Ollama server.

    Returns the raw response string (expected to be JSON) or ``None`` on failure.
    A very low temperature (0.05) is used to get deterministic JSON output.
    Guarded by a configurable timeout (``RICE_LLM_TIMEOUT`` env var, default 90s).
    """
    if not _LLM_OK:
        return None
    # Sanitise message content
    safe_messages = [
        {**m, "content": _escape_prompt(m.get("content", ""))} for m in messages
    ]
    try:
        client = _ollama.Client(host=_OLLAMA_HOST)
        response = client.chat(
            model=_LLM_MODEL,
            messages=safe_messages,
            format="json",
            options={"temperature": temperature, "num_predict": 2048, "num_gpu": 0},
            timeout=_LLM_TIMEOUT,
        )
        return response["message"]["content"]
    except Exception as exc:
        logger.warning("LLM call failed (%s@%s): %s", _LLM_MODEL, _OLLAMA_HOST, exc)
        return None


def _llm_extract_metadata(text: str) -> Dict[str, Any]:
    """Use LLM to extract fiche module metadata (NER).

    Returns a partial dict with only the fields the LLM found.
    The caller merges this with the table/regex NER results.
    """
    import json as _json_local
    truncated = text[:3500]
    prompt = (
        "Tu es un expert en extraction d'information de fiches modules universitaires françaises.\n"
        "Extrais les métadonnées de cette fiche module ESPRIT (école d'ingénieurs tunisienne).\n"
        "Retourne UNIQUEMENT un objet JSON valide avec ces clés (null si absent) :\n"
        "{\n"
        '  "code_module": "code alphanumérique du module (ex: MT-34, GC05-F)",\n'
        '  "nom_module": "intitulé complet du module/UE",\n'
        '  "unite_pedagogique": "nom de l\'unité pédagogique ou spécialité",\n'
        '  "responsable": "nom complet du responsable/coordinateur du module",\n'
        '  "enseignants_noms": ["liste", "des", "noms", "complets"],\n'
        '  "enseignants_roles": {"Prénom NOM": "responsable|coordinateur|enseignant"},\n'
        '  "prerequis": "prérequis du module (chaîne texte)",\n'
        '  "objectif": "objectifs pédagogiques du module (chaîne texte)"\n'
        "}\n\n"
        f"FICHE :\n{truncated}"
    )
    raw = _llm_chat([{"role": "user", "content": prompt}])
    if not raw:
        return {}
    try:
        result = _json_local.loads(raw)
        out: Dict[str, Any] = {}
        if result.get("code_module") and isinstance(result["code_module"], str):
            code = result["code_module"].strip().upper()
            if re.search(r"[A-Z0-9]", code) and len(code) >= 2:
                out["code_module"] = code
        if result.get("nom_module") and isinstance(result["nom_module"], str):
            nom = result["nom_module"].strip()
            if len(nom) > 3:
                out["nom_module"] = nom
        if result.get("unite_pedagogique") and isinstance(result["unite_pedagogique"], str):
            up = result["unite_pedagogique"].strip()
            if len(up) > 3:
                out["unite_pedagogique"] = up
        if result.get("responsable") and isinstance(result["responsable"], str):
            resp = _clean_name(result["responsable"])
            if resp:
                out["responsable"] = resp
        if result.get("enseignants_noms") and isinstance(result["enseignants_noms"], list):
            names = [_clean_name(str(n)) for n in result["enseignants_noms"] if n]
            names = [n for n in names if n]
            if names:
                out["enseignants_noms"] = names
        if result.get("enseignants_roles") and isinstance(result["enseignants_roles"], dict):
            valid_roles = {"responsable", "coordinateur", "enseignant", "intervenant"}
            roles: Dict[str, str] = {}
            for k, v in result["enseignants_roles"].items():
                ck = _clean_name(str(k))
                if ck and str(v) in valid_roles:
                    roles[ck] = str(v)
            if roles:
                out["enseignants_roles"] = roles
        if result.get("prerequis") and isinstance(result["prerequis"], str):
            pq = result["prerequis"].strip()
            if len(pq) > 3:
                out["prerequis"] = pq
        if result.get("objectif") and isinstance(result["objectif"], str):
            obj = result["objectif"].strip()
            if len(obj) > 5:
                out["objectif"] = obj
        logger.info(f"LLM metadata extracted: {list(out.keys())}")
        return out
    except Exception as exc:
        logger.warning(f"LLM metadata parse error: {exc} | raw={raw[:200]}")
        return {}


def _llm_extract_acquis(text: str) -> List[Dict[str, Any]]:
    """Use LLM to extract Acquis d'Apprentissage (AA) with Bloom levels.

    Returns a list of ``{id, text, bloom_level}`` dicts.
    Empty list on failure (caller falls back to regex parser).
    """
    import json as _json_local
    aa_match = re.search(
        r"(Acquis\s+d['\u2019\u2018]apprentissage.*?)(?=Contenu\s+d[eé]taill[eé]|Plan\s+du\s+cours|$)",
        text, re.I | re.DOTALL,
    )
    section = (aa_match.group(1) if aa_match else text)[:2800]
    prompt = (
        "Tu es un expert en ingénierie pédagogique. "
        "Extrais les Acquis d'Apprentissage (AA) de cette section d'une fiche module universitaire française.\n"
        "Chaque AA est en général introduit par 'AA1', 'AA2', etc.\n"
        "Pour chaque AA, attribue un niveau de la taxonomie de Bloom :\n"
        "  1=Mémoriser, 2=Comprendre, 3=Appliquer, 4=Analyser, 5=Évaluer, 6=Créer\n"
        "Retourne UNIQUEMENT un objet JSON valide :\n"
        '{"acquis": [{"id": 1, "text": "Texte complet de l\'AA", "bloom_level": 3}, ...]}\n\n'
        f"SECTION :\n{section}"
    )
    raw = _llm_chat([{"role": "user", "content": prompt}])
    if not raw:
        return []
    try:
        result = _json_local.loads(raw)
        validated = []
        for aa in result.get("acquis", []):
            t = str(aa.get("text", "")).strip()
            bl = int(aa.get("bloom_level", 2))
            if len(t) > 5:
                validated.append({
                    "id": int(aa.get("id", len(validated) + 1)),
                    "text": t,
                    "bloom_level": min(max(bl, 1), 6),
                })
        logger.info(f"LLM acquis extracted: {len(validated)}")
        return validated
    except Exception as exc:
        logger.warning(f"LLM acquis parse error: {exc} | raw={raw[:200]}")
        return []


def _llm_extract_seances(text: str) -> List[Dict[str, Any]]:
    """Use LLM to extract sessions/séances from the contenu détaillé.

    Returns a list of séance dicts compatible with ``_extract_seances()``.
    Empty list on failure.
    """
    import json as _json_local
    seance_match = re.search(
        r"(Contenu\s+d[eé]taill[eé].*?)(?=Mode\s+d['\u2019]|[EÉ]valuation|R[eé]f[eé]rences|Bibliographie|$)",
        text, re.I | re.DOTALL,
    )
    section = (seance_match.group(1) if seance_match else text)[:3000]
    prompt = (
        "Tu es un expert en ingénierie pédagogique. "
        "Extrais les séances/chapitres/sessions du contenu détaillé de cette fiche module universitaire.\n"
        "Pour chaque séance, identifie : numéro, titre, sous-points, type (Cours/TP/TD/Projet), durée.\n"
        "Retourne UNIQUEMENT un objet JSON valide :\n"
        '{"seances": [{\n'
        '  "numero": "1",\n'
        '  "titre": "Titre de la séance",\n'
        '  "items": ["Sous-point A", "Sous-point B"],\n'
        '  "type_apprentissage": "Cours|TP|TD|Projet|null",\n'
        '  "duree": "3h|null"\n'
        '}]}\n\n'
        f"SECTION :\n{section}"
    )
    raw = _llm_chat([{"role": "user", "content": prompt}])
    if not raw:
        return []
    try:
        result = _json_local.loads(raw)
        validated = []
        for s in result.get("seances", []):
            titre = str(s.get("titre", "")).strip()
            if not titre:
                continue
            validated.append({
                "numero": str(s.get("numero", len(validated) + 1)),
                "titre": titre,
                "items": [str(i).strip() for i in s.get("items", []) if str(i).strip()],
                "type_apprentissage": s.get("type_apprentissage") or None,
                "duree": s.get("duree") or None,
            })
        logger.info(f"LLM séances extracted: {len(validated)}")
        return validated
    except Exception as exc:
        logger.warning(f"LLM séances parse error: {exc} | raw={raw[:200]}")
        return []


def _llm_fallback_items(text: str, module_name: str) -> List[str]:
    """Use LLM to extract competence items from freeform text (last-resort fallback).

    Returns a list of action-verb phrases suitable as savoir names.
    """
    import json as _json_local
    prompt = (
        "Tu es un expert en ingénierie pédagogique. "
        "Extrais les compétences et savoirs à acquérir de ce texte de fiche module.\n"
        "Formule chaque compétence comme une phrase courte avec un verbe d'action "
        "(taxonomie de Bloom : analyser, concevoir, appliquer, réaliser…).\n"
        "Retourne UNIQUEMENT un objet JSON valide :\n"
        '{"items": ["Analyser les données de sol", "Concevoir un coffrage", ...]}\n\n'
        f"MODULE : {module_name}\n\nTEXTE :\n{text[:2500]}"
    )
    raw = _llm_chat([{"role": "user", "content": prompt}])
    if not raw:
        return []
    try:
        result = _json_local.loads(raw)
        items = [str(i).strip() for i in result.get("items", []) if len(str(i).strip()) > 5]
        logger.info(f"LLM fallback items extracted: {len(items)}")
        return items[:25]
    except Exception as exc:
        logger.warning(f"LLM fallback parse error: {exc} | raw={raw[:200]}")
        return []


# ─────────────────────────────────────────────────────────────────────────────
# NLP – Sous-compétence title extraction (regex + LLM fallback)
# ─────────────────────────────────────────────────────────────────────────────

# Various forms seen in ESPRIT fiches:
#   "SC1 – Analyse des exigences"
#   "Sous-compétence : Conception du test"
#   "S‑C 2. Validation logicielle"
_RE_SUBCOMP_TITLE_1 = re.compile(
    r"^(?:SC|Sous[-\s]?comp[ée]tence)\s*\d*\s*[:\-–]\s*([^\n\r]+)", re.I | re.M
)
_RE_SUBCOMP_TITLE_2 = re.compile(
    r"^Sous[-\s]?comp[ée]tence\s*[:\-–]\s*([^\n\r]+)", re.I | re.M
)
_RE_SUBCOMP_TITLE_3 = re.compile(
    r"^S[\-‑]C\s*(\d+)\s*[.\-–]\s*([^\n\r]+)", re.I | re.M
)


def _extract_subcompetences(text: str) -> List[Tuple[int, str]]:
    """Extract explicit sous-compétence titles from fiche text.

    Returns a list of ``(line_number_1based, title_string)`` tuples, sorted by
    line number.  Returns an empty list when no explicit titles are found in
    the text (the caller should then fall back to the LLM extractor or the
    default Bloom-level grouping).
    """
    titles: List[Tuple[int, str]] = []
    seen_titles: set = set()
    lines = text.splitlines()
    for i, line in enumerate(lines):
        for pat in (_RE_SUBCOMP_TITLE_1, _RE_SUBCOMP_TITLE_2, _RE_SUBCOMP_TITLE_3):
            m = pat.search(line)
            if m:
                # _RE_SUBCOMP_TITLE_3 captures (number, title), others capture (title)
                titre = m.group(2) if len(m.groups()) >= 2 else m.group(1)
                titre = titre.strip(" :‑–-")
                if titre and titre.lower() not in seen_titles:
                    seen_titles.add(titre.lower())
                    titles.append((i + 1, titre))
                break
    return titles


def _llm_extract_subcompetences(text: str, module_name: str) -> List[str]:
    """Use LLM to extract sous-compétence titles when regex finds nothing.

    Returns a list of title strings (no line numbers).
    """
    prompt = (
        "Tu es un expert en ingénierie pédagogique. "
        "À partir du texte ci-dessous, liste seulement les titres de sous-compétences, "
        "un titre par ligne. Ne numérote pas les titres. Pas de commentaire.\n\n"
        f"MODULE : {module_name}\n\n"
        f"{text[:3500]}"
    )
    raw = _llm_chat([{"role": "user", "content": prompt}], temperature=0.1)
    if not raw:
        return []
    # LLM may return JSON or plain lines; handle both
    try:
        import json as _jl
        parsed = _jl.loads(raw)
        if isinstance(parsed, list):
            return [str(t).strip() for t in parsed if str(t).strip()]
        if isinstance(parsed, dict):
            items = parsed.get("items") or parsed.get("titres") or parsed.get("sous_competences") or []
            return [str(t).strip() for t in items if str(t).strip()]
    except Exception:
        pass
    return [ln.strip() for ln in raw.splitlines() if ln.strip() and len(ln.strip()) > 3]


# ─────────────────────────────────────────────────────────────────────────────
# NLP – Fiche metadata extraction (Named Entity Recognition style)
# ─────────────────────────────────────────────────────────────────────────────

# ── Standard format: label : value on the SAME line ─────────────────────────
_RE_MODULE_CODE = re.compile(
    r"(?:Code|code)\s*[:\-]?\s*([A-Z][A-Z0-9\-_]{2,15})", re.I
)
# ── Table/reversed format: value on previous line, label on next line ────────
# Captures code like "MT-34" that appears as a standalone token on its own line
_RE_MODULE_CODE_TABLE = re.compile(
    r"^\s*([A-Z]{1,4}[\-_]?\d{1,4}[A-Z]?)\s*\d*h.*$",
    re.MULTILINE,
)
_RE_MODULE_NAME = re.compile(
    r"(?:Module|Mati\u00e8re|Unit\u00e9\s+d['\u2019]enseignement)\s*[:\-]?\s*(.{5,80})",
    re.IGNORECASE,
)
_RE_UNITE_PEDAGOGIQUE = re.compile(
    r"(?:Unit\u00e9\s+p\u00e9dagogique|UP)\s*[:\-]?\s*(.{3,60})",
    re.IGNORECASE,
)
# Standard responsable (label: value)
_RE_RESPONSABLE = re.compile(
    r"(?:Responsable(?:\s+(?:Module|UE|Mati[eè]re|Cours))?|Coordinat(?:eur|rice))\s*[:\-]?\s*(.{3,80})",
    re.IGNORECASE,
)
# Reversed format: capture line BEFORE "Responsable Module" label
_RE_RESPONSABLE_REV = re.compile(
    r"^(.{5,80})\n\s*Responsable(?:\s+(?:Module|UE|Mati[eè]re|Cours))",
    re.IGNORECASE | re.MULTILINE,
)
# Standard enseignants (label: value)
_RE_ENSEIGNANTS = re.compile(
    r"(?:Enseignants?|Intervenants?|Enseignants?\s*[\u2013\-]\s*Intervenants?|Professeurs?|Formateurs?|Titulaire)\s*[:\-]?\s*(.{5,300})",
    re.IGNORECASE,
)
# Reversed format: capture line BEFORE "Enseignants" label
_RE_ENSEIGNANTS_REV = re.compile(
    r"^(.{5,300})\n\s*(?:Enseignants?(?:\s*[\u2013\-]\s*Intervenants?)?|Intervenants?)\s*$",
    re.IGNORECASE | re.MULTILINE,
)
# Additional patterns for ESPRIT fiche modules (table-based PDFs)
_RE_NOM_PRENOM = re.compile(
    r"(?:Nom\s+(?:et|&)\s+Pr[eé]nom|Nom\s+Pr[eé]nom)\s*[:\-]?\s*(.{5,80})",
    re.IGNORECASE,
)
_RE_EQUIPE_PEDAGOGIQUE = re.compile(
    r"[EÉ]quipe\s+[Pp][eé]dagogique\s*[:\-]?\s*(.{5,400})",
    re.IGNORECASE,
)
_RE_COORDINATEUR = re.compile(
    r"Coordinat(?:eur|rice)(?:\s+(?:du\s+)?(?:module|UE|cours))?\s*[:\-]?\s*(.{3,80})",
    re.IGNORECASE,
)
# Standard prerequis
_RE_PREREQUIS = re.compile(
    r"(?:Pr\u00e9requis|Pr\u00e9[\-\s]?requis)\s*[:\-]?\s*(.{3,200})",
    re.IGNORECASE,
)
# Reversed format: capture line BEFORE "Prérequis" label
_RE_PREREQUIS_REV = re.compile(
    r"^(.{3,200})\n\s*(?:Pr[eé]requis|Pr[eé][\-\s]?requis)\s*$",
    re.IGNORECASE | re.MULTILINE,
)
_RE_OBJECTIF = re.compile(
    r"(?:Objectif(?:s)?(?:\s+du\s+module)?)\s*[:\-]?\s*(.+?)(?=\n\n|Mode\s+d|Acquis|$)",
    re.IGNORECASE | re.DOTALL,
)

# ─── Name cleaning & splitting helpers ───────────────────────────────────────

# Words that are NOT person names (false positive filters)
_STOP_WORDS = {
    "module", "cours", "matiere", "matière", "ue", "tp", "td", "prerequis",
    "objectif", "objectifs", "contenu", "evaluation", "mode", "duree", "durée",
    "semestre", "niveau", "credits", "coefficient", "code", "reference",
    "département", "departement", "informatique", "genie", "civil", "esprit",
    "fiche", "pedagogique", "pédagogique", "unite", "unité", "formation",
    "description", "compétence", "competence", "savoir", "acquis",
    "apprentissage", "séance", "seance", "chapitre",
    "responsable", "coordinateur", "coordinatrice", "enseignant", "enseignants",
    "intervenant", "intervenants", "web", "semantique", "sémantique",
    "nouvelles", "applications", "options", "niveaux",
    # Fiche module table headers & labels
    "he", "hne", "ects", "integre", "intégré", "detaille", "détaillé",
    "situation", "rendu", "rendus", "atelier", "projet",
    "derniere", "dernière", "mise", "jour", "date",
    "moyenne", "calculee", "calculée", "suivant", "suit",
}

def _clean_name(raw: str) -> Optional[str]:
    """Clean a potential person name: remove noise, validate it looks like a name."""
    # Remove common trailing noise
    name = re.sub(r"\s*[\(\[].*?[\)\]]", "", raw)  # remove (parentheses)
    name = re.sub(r"\s*[-–]\s*(cours|tp|td|module|ue).*$", "", name, flags=re.I)
    name = re.sub(r"\s+(mail|email|tél|tel|bureau|grade).*$", "", name, flags=re.I)
    # Strip academic titles
    name = re.sub(r"^(Dr\.?|Pr\.?|Prof\.?|M\.?|Mme\.?|Mr\.?)\s+", "", name, flags=re.I)
    name = name.strip().strip(".,;:-–")
    # Take only first line if multiline
    name = name.split("\n")[0].strip()
    # Strip embedded enseignant/module codes (e.g. "GC05", "E001") before digit check
    # so a string like "GC05 Abidi Mounir" is not thrown away entirely
    name_no_codes = re.sub(r'\b[A-Z]{1,5}\d{1,4}\b\s*', '', name).strip()
    if name_no_codes:   # only use stripped version if something remains
        name = name_no_codes.strip(".,;:- –")
    # Must have at least 2 words (first + last name)
    words = name.split()
    if len(words) < 2 or len(name) < 5:
        return None
    # Filter out if it's all stop words
    norm_words = [_normalize(w) for w in words]
    if all(w in _STOP_WORDS for w in norm_words):
        return None
    # Reject if it contains digits (not a name — after codes are stripped)
    if re.search(r"\d", name):
        return None
    # Reject if ALL-CAPS abbreviation (like "HNE ECTS", "HE HNE")
    if all(w.isupper() and len(w) <= 4 for w in words):
        return None
    return name.strip()

def _split_names(raw: str) -> List[str]:
    """Split a raw string of professor names separated by , ; / – - or newlines."""
    # First split by common separators (including dash/en-dash which ESPRIT uses)
    parts = re.split(r"[,;/\n]+", raw)
    # Also try " – " (en-dash with spaces) and " - " as separators
    expanded = []
    for p in parts:
        sub = re.split(r"\s+[–\-]\s+", p)
        expanded.extend(sub)
    # Also try " et " as separator
    final = []
    for p in expanded:
        sub = re.split(r"\s+et\s+", p, flags=re.IGNORECASE)
        final.extend(sub)
    names = []
    for part in final:
        cleaned = _clean_name(part.strip())
        if cleaned:
            names.append(cleaned)
    return names

# ── Labels used for table-based NER (normalised) ─────────────────────────────
_TABLE_NER_LABELS: Dict[str, str] = {
    # normalised label text → meta key
    "responsable":             "responsable",
    "responsable du module":   "responsable",
    "responsable module":      "responsable",
    "coordinateur":            "coordinateur",
    "coordinatrice":           "coordinateur",
    "coordinateur du module":  "coordinateur",
    "enseignant":              "enseignant_raw",
    "enseignants":             "enseignant_raw",
    "intervenants":            "enseignant_raw",
    "equipe pedagogique":      "enseignant_raw",
    "nom et prenom":           "enseignant_raw",
    "nom prenom":              "enseignant_raw",
    "intitule":                "nom_module",
    "intitule du module":      "nom_module",
    "module":                  "nom_module",
    "code":                    "code_module",
    "code module":             "code_module",
    "code ue":                 "code_module",
    "unite pedagogique":       "unite_pedagogique",
    "up":                      "unite_pedagogique",
    "prerequis":               "prerequis",
    "pre-requis":              "prerequis",
    "objectif":                "objectif",
    "objectifs":               "objectif",
    "objectifs du module":     "objectif",
}


def _scan_tables_for_meta(raw_tables: List, meta: Dict[str, Any]) -> None:
    """Scan pdfplumber raw tables to extract metadata fields.

    Each table is a list-of-rows; each row is a list of cell strings.
    We look for a cell whose normalised text matches a known label and
    take the adjacent cell (right or below) as the value.

    Results are written into *meta* only for keys that are not yet set,
    so this function acts as the **primary** source; regex is the fallback.
    """
    for table in raw_tables:
        for row_idx, row in enumerate(table):
            cells = [str(c).strip() if c else "" for c in row]
            for col_idx, cell in enumerate(cells):
                norm_cell = _normalize(cell)
                meta_key = _TABLE_NER_LABELS.get(norm_cell)
                if meta_key is None:
                    continue
                # Try to get a value: right neighbour first, then cell below
                value = ""
                if col_idx + 1 < len(cells) and cells[col_idx + 1].strip():
                    value = cells[col_idx + 1].strip()
                elif row_idx + 1 < len(table):
                    next_row = [str(c).strip() if c else "" for c in table[row_idx + 1]]
                    if col_idx < len(next_row) and next_row[col_idx].strip():
                        value = next_row[col_idx].strip()
                if not value:
                    continue

                if meta_key == "responsable" and "responsable" not in meta:
                    cleaned = _clean_name(value)
                    if cleaned:
                        meta["responsable"] = cleaned

                elif meta_key == "coordinateur" and "coordinateur" not in meta:
                    cleaned = _clean_name(value)
                    if cleaned:
                        meta["coordinateur"] = cleaned
                        meta.setdefault("enseignants_noms", []).append(cleaned)
                        meta.setdefault("enseignants_roles", {})[cleaned] = "coordinateur"

                elif meta_key == "enseignant_raw":
                    names = _split_names(value)
                    for n in names:
                        meta.setdefault("enseignants_noms", []).append(n)
                        meta.setdefault("enseignants_roles", {})[n] = "enseignant"

                elif meta_key == "nom_module" and "nom_module" not in meta:
                    raw = value.strip().rstrip(".")
                    raw = re.sub(r"\s*(Pr\u00e9requis|Niveaux|Objectif|Derni[\u00e8e]re).*$",
                                 "", raw, flags=re.I)
                    if len(raw) > 2:
                        meta["nom_module"] = raw

                elif meta_key == "code_module" and "code_module" not in meta:
                    code = value.strip().upper()
                    if re.search(r"\d", code):
                        meta["code_module"] = code

                elif meta_key == "unite_pedagogique" and "unite_pedagogique" not in meta:
                    if len(value) > 2:
                        meta["unite_pedagogique"] = value.strip()

                elif meta_key == "prerequis" and "prerequis" not in meta:
                    meta["prerequis"] = value.strip()

                elif meta_key == "objectif" and "objectif" not in meta:
                    meta["objectif"] = re.sub(r"\s*\n\s*", " ", value).strip()


def _extract_metadata(text: str, raw_tables: Optional[List] = None) -> Dict[str, Any]:
    """
    NLP-based Named Entity extraction for fiche module metadata.

    Strategy (priority order):
    1. **LLM-based NER** (primary) — Ollama local model understands any layout.
    2. **Table-based NER** (fills gaps) — pdfplumber structured tables.
    3. **Regex-based NER** (fills remaining gaps) — pattern matching fallback.
    """
    meta: Dict[str, Any] = {}

    # ── 0. LLM-based NER (primary, highest precision) ─────────────────────
    if _LLM_OK:
        llm_meta = _llm_extract_metadata(text)
        # Seed meta with LLM findings – only authoritative non-empty values
        for key in ("code_module", "nom_module", "unite_pedagogique",
                    "responsable", "prerequis", "objectif"):
            if llm_meta.get(key):
                meta[key] = llm_meta[key]
        if llm_meta.get("enseignants_noms"):
            meta["enseignants_noms"] = llm_meta["enseignants_noms"]
        if llm_meta.get("enseignants_roles"):
            meta["enseignants_roles"] = llm_meta["enseignants_roles"]

    # ── 1. Table-based NER (fills any gaps not found by LLM) ──────────────
    if raw_tables:
        _scan_tables_for_meta(raw_tables, meta)

    # ── 2. Regex-based NER (fallback for any field not yet populated) ─────

    # ── Code module ──────────────────────────────────────────────────────
    # Table format first (more specific): "MT-34 24h 30h 3" line
    m = _RE_MODULE_CODE_TABLE.search(text)
    if m:
        meta["code_module"] = m.group(1).strip().upper()
    # Standard format fallback (label: value) — require at least one digit
    if "code_module" not in meta:
        m = _RE_MODULE_CODE.search(text)
        if m:
            code = m.group(1).strip().upper()
            if re.search(r'\d', code):
                meta["code_module"] = code

    # ── Module name ──────────────────────────────────────────────────────
    if "nom_module" not in meta:
        m = _RE_MODULE_NAME.search(text)
        if m:
            name = m.group(1).strip().rstrip(".")
            name = re.sub(r"\s*(Pr\u00e9requis|Niveaux|Objectif|Derni[eè]re).*$", "", name, flags=re.I)
            meta["nom_module"] = name

    # ── Unité pédagogique ────────────────────────────────────────────────
    m = _RE_UNITE_PEDAGOGIQUE.search(text)
    if m:
        val = m.group(1).strip()
        # In table format, "Unité pédagogique" is a label and value is on the line above
        # Check if the captured value looks like a label (e.g. starts with "Nouvelles" — keep it)
        meta["unite_pedagogique"] = val
    # Also try: line BEFORE "Unité pédagogique"
    if "unite_pedagogique" not in meta:
        m_rev = re.search(
            r"^(.{3,60})\n\s*(?:Unit[eé]\s+p[eé]dagogique|UP)\s*$",
            text, re.I | re.MULTILINE
        )
        if m_rev:
            meta["unite_pedagogique"] = m_rev.group(1).strip()

    # ── Responsable ──────────────────────────────────────────────────────
    # Try reversed format first (value line BEFORE label)
    m_rev = _RE_RESPONSABLE_REV.search(text)
    if m_rev:
        raw_resp = m_rev.group(1).strip()
        # Validate it looks like a person name (not a section header)
        cleaned = _clean_name(raw_resp)
        if cleaned:
            meta["responsable"] = cleaned
    # Standard format fallback
    if "responsable" not in meta:
        m = _RE_RESPONSABLE.search(text)
        if m:
            raw = m.group(1).strip()
            # If it looks like a label/section rather than a name, skip
            cleaned = _clean_name(raw)
            if cleaned:
                meta["responsable"] = cleaned

    # ── Enseignants ──────────────────────────────────────────────────────
    ens_names: List[str] = []
    # Try reversed format first (names line BEFORE "Enseignants" label)
    m_rev = _RE_ENSEIGNANTS_REV.search(text)
    if m_rev:
        raw = m_rev.group(1).strip()
        ens_names = _split_names(raw)
    # Standard format fallback
    if not ens_names:
        m = _RE_ENSEIGNANTS.search(text)
        if m:
            raw = m.group(1).strip()
            ens_names = _split_names(raw)
    if ens_names:
        meta["enseignants_noms"] = ens_names
        # Assign roles
        for n in ens_names:
            meta.setdefault("enseignants_roles", {})[n] = "enseignant"

    # Additional ESPRIT-specific patterns
    m = _RE_NOM_PRENOM.search(text)
    if m:
        name = _clean_name(m.group(1).strip())
        if name:
            meta.setdefault("enseignants_noms", []).append(name)
            meta.setdefault("enseignants_roles", {})[name] = "enseignant"

    m = _RE_EQUIPE_PEDAGOGIQUE.search(text)
    if m:
        raw = m.group(1).strip()
        eq_names = _split_names(raw)
        for n in eq_names:
            meta.setdefault("enseignants_noms", []).append(n)
            meta.setdefault("enseignants_roles", {})[n] = "enseignant"

    m = _RE_COORDINATEUR.search(text)
    if m:
        name = _clean_name(m.group(1).strip())
        if name:
            meta.setdefault("enseignants_noms", []).append(name)
            meta.setdefault("enseignants_roles", {})[name] = "coordinateur"

    # Assign role to responsable
    if "responsable" in meta:
        meta.setdefault("enseignants_roles", {})[meta["responsable"]] = "responsable"

    if "responsable" in meta and "enseignants_noms" not in meta:
        meta["enseignants_noms"] = [meta["responsable"]]

    # Deduplicate names
    if "enseignants_noms" in meta:
        seen = set()
        unique = []
        for n in meta["enseignants_noms"]:
            norm = _normalize(n)
            if norm not in seen and len(n.strip()) > 3:
                seen.add(norm)
                unique.append(n)
        meta["enseignants_noms"] = unique

    # ── Prérequis ────────────────────────────────────────────────────────
    # Try reversed format first (value BEFORE label)
    m_rev = _RE_PREREQUIS_REV.search(text)
    if m_rev:
        meta["prerequis"] = m_rev.group(1).strip()
    if "prerequis" not in meta:
        m = _RE_PREREQUIS.search(text)
        if m:
            val = m.group(1).strip()
            # If in table format, the captured text may be the next label — check
            if not re.match(r"^\d+[A-Z]{3,}", val):  # not "5TWIN" (that's a level)
                meta["prerequis"] = val

    # ── Objectif ─────────────────────────────────────────────────────────
    m = _RE_OBJECTIF.search(text)
    if m:
        obj = m.group(1).strip()
        # Clean up multi-line wrapping
        obj = re.sub(r"\s*\n\s*", " ", obj).strip()
        meta["objectif"] = obj

    return meta

# ─────────────────────────────────────────────────────────────────────────────
# NLP – Acquis d'Apprentissage extraction (AA)
# ─────────────────────────────────────────────────────────────────────────────

# ── AA extraction patterns ───────────────────────────────────────────────────
# Standard: AA1 <text> <bloom_level>  (all on one line)
_RE_AA_LINE = re.compile(
    r"(?:AA\s*(\d+))\s+(.+?)\s+(\d)\s*$", re.MULTILINE
)
# Alternative: AA lines without explicit level
_RE_AA_ALT = re.compile(
    r"(?:AA\s*(\d+))\s+(.+)", re.MULTILINE
)


def _extract_acquis_apprentissage(text: str) -> List[Dict[str, Any]]:
    """
    NLP extraction of Acquis d'Apprentissage with Bloom levels.

    Strategy:
    1. **LLM** (primary) — understands any AA format, assigns Bloom levels semantically.
    2. **Regex block-parser** (fallback) — line-by-line AA marker detection.
    Returns list of {id, text, bloom_level}.
    """
    # ── 0. LLM extraction (primary) ──────────────────────────────────────
    if _LLM_OK:
        llm_acquis = _llm_extract_acquis(text)
        if llm_acquis:
            return llm_acquis

    # ── 1. Regex block parser (fallback) ─────────────────────────────────
    acquis = []

    # ── Find the AA block section ─────────────────────────────────────────
    aa_block_match = re.search(
        r"Acquis\s+d['’]apprentissage\s*:?\s*.*?\n(.+?)(?=Contenu\s+d[eé]taill[eé]|Plan\s+du\s+cours|$)",
        text, re.I | re.DOTALL
    )
    if not aa_block_match:
        return []

    block = aa_block_match.group(1)
    lines = block.split('\n')

    # ── Step 1: Parse lines into markers and text segments ────────────────
    parsed = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Skip header, footer, and legend lines
        if re.match(r'^AA\s+Acquis|^Niveau|^\*\s*:|^\(1\s*:', stripped, re.I):
            continue
        if stripped.startswith('*') or stripped.startswith('(1'):
            continue
        # Skip legend continuation (e.g. ": Créer)")
        if re.match(r'^:\s', stripped):
            continue

        m = re.match(r'^AA\s*(\d+)\s*(.*)', stripped)
        if m:
            parsed.append({'type': 'marker', 'aa': int(m.group(1)), 'rest': m.group(2).strip()})
        else:
            parsed.append({'type': 'text', 'content': stripped})

    # ── Step 2: Build AA segments ─────────────────────────────────────────
    # Between consecutive AA markers, text lines are split into:
    # - post-text (continuation of previous AA): starts with lowercase
    # - pre-text (start of next AA): starts with uppercase verb/noun
    marker_indices = [i for i, p in enumerate(parsed) if p['type'] == 'marker']

    if not marker_indices:
        # Fallback: try simple single-line regex
        for m in _RE_AA_LINE.finditer(text):
            acquis.append({
                "id": int(m.group(1)),
                "text": m.group(2).strip(),
                "bloom_level": min(max(int(m.group(3)), 1), 6),
            })
        return acquis

    segments = []
    for idx, mi in enumerate(marker_indices):
        aa_num = parsed[mi]['aa']
        rest = parsed[mi]['rest']

        # Extract bloom level from end of inline text
        bloom = 0
        bm = re.search(r'\s+(\d)\s*$', rest)
        if bm:
            bloom = int(bm.group(1))
            rest = rest[:bm.start()].strip()
        elif re.match(r'^(\d)\s*$', rest):
            bloom = int(rest)
            rest = ''

        # Collect text lines between previous marker and this marker
        prev_mi = marker_indices[idx - 1] if idx > 0 else -1
        text_between = []
        for j in range(prev_mi + 1, mi):
            if parsed[j]['type'] == 'text':
                text_between.append(parsed[j]['content'])

        # Split text_between into post-text (prev AA) and pre-text (this AA)
        # Heuristic: first line starting with uppercase (new sentence) marks the split
        split_point = len(text_between)  # default: all are continuation (post-text)
        for k, t in enumerate(text_between):
            if not t:
                continue
            first_char = t[0]
            if first_char.isupper():
                # Exclude common continuation words
                if not re.match(r'^(sur|de|du|des|le|la|les|un|une|et|ou|au|aux)\s', t, re.I):
                    split_point = k
                    break

        post_text_prev = text_between[:split_point]
        pre_text_this = text_between[split_point:]

        # Assign post-text to previous segment
        if post_text_prev and segments:
            segments[-1]['post'].extend(post_text_prev)

        # Collect trailing text for the last marker
        trailing = []
        if idx == len(marker_indices) - 1:
            for j in range(mi + 1, len(parsed)):
                if parsed[j]['type'] == 'text':
                    trailing.append(parsed[j]['content'])

        segments.append({
            'aa': aa_num,
            'bloom': bloom,
            'pre': list(pre_text_this),
            'inline': rest,
            'post': trailing,
        })

    # ── Step 3: Build final result ────────────────────────────────────────
    for seg in segments:
        parts = seg['pre'] + ([seg['inline']] if seg['inline'] else []) + seg['post']
        aa_text = ' '.join(parts).strip()
        bloom = seg['bloom']

        # Clean trailing noise (table labels)
        aa_text = re.sub(
            r"\s*(?:Situation|Dur[eé]e|Rendu|d['’]apprentissage).*$",
            "", aa_text, flags=re.I
        ).strip()

        if not bloom and aa_text:
            bloom = _detect_bloom_level(aa_text)

        if aa_text and len(aa_text) > 5:
            acquis.append({
                "id": seg['aa'],
                "text": aa_text,
                "bloom_level": min(max(bloom, 1), 6),
            })

    return acquis

# ─────────────────────────────────────────────────────────────────────────────
# NLP – Séance / Session extraction (Contenu détaillé)
# ─────────────────────────────────────────────────────────────────────────────

_RE_SEANCE = re.compile(
    r"(?:S[eé]ance|Seance|Session|Chapitre|Semaine)\s*(\d+(?:\s*[-\u2013]\s*\d+)?)\s*[:\-]?\s*(.+?)(?=\n)",
    re.IGNORECASE,
)
_RE_CHECKMARK = re.compile(r"^[\u2714\u2713\u2611\u2610]\s*(.+)$", re.MULTILINE)
_RE_BULLET    = re.compile(r"^[\-\u2022\*\u203A\u25E6\u25AA]\s+(.+)$", re.MULTILINE)
_RE_NUMBERED  = re.compile(r"^\d+[\.\)]\s+(.+)$", re.MULTILINE)

def _extract_seances(text: str) -> List[Dict[str, Any]]:
    """
    NLP extraction of séances/sessions from the contenu détaillé.

    Strategy:
    1. **LLM** (primary) — handles any section format and bullet style.
    2. **Regex** (fallback) — requires explicit "Séance N" markers.
    """
    # ── 0. LLM extraction (primary) ──────────────────────────────────────
    if _LLM_OK:
        llm_seances = _llm_extract_seances(text)
        if llm_seances:
            return llm_seances

    # ── 1. Regex fallback ─────────────────────────────────────────────────
    seances = []
    matches = list(_RE_SEANCE.finditer(text))
    if not matches:
        return seances

    for i, m in enumerate(matches):
        numero = m.group(1).strip()
        titre = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block = text[start:end]

        items = []
        for pattern in [_RE_CHECKMARK, _RE_BULLET, _RE_NUMBERED]:
            items.extend(pattern.findall(block))
        items = [it.strip() for it in items if len(it.strip()) > 5]

        type_match = re.search(
            r"(?:Situation\s*(?:\(s\))?\s*|Type\s*)[:\-]?\s*(cours\s+int[e\u00e9]gr[e\u00e9]|TP|TD|APP|Projet|Labo)",
            block, re.IGNORECASE,
        )
        type_apprentissage = type_match.group(1).strip() if type_match else None

        duree_match = re.search(r"(?:Dur\u00e9e|Duree)\s*[:\-]?\s*(\d+\s*h)", block, re.I)
        duree = duree_match.group(1).strip() if duree_match else None

        seances.append({
            "numero": numero,
            "titre": titre,
            "items": items,
            "type_apprentissage": type_apprentissage,
            "duree": duree,
        })

    return seances

# ─────────────────────────────────────────────────────────────────────────────
# NLP – Enseignant matching (fuzzy name matching)
# ─────────────────────────────────────────────────────────────────────────────

# Minimum fuzzy-match score (0-100) to accept a name match
_FUZZY_THRESHOLD = 82


def _match_enseignants_by_name(
    fiche_names: List[str],
    enseignants: List[EnseignantInfo],
) -> Tuple[List[str], Dict[str, Tuple[str, str]]]:
    """Fuzzy match enseignant names found in the fiche against the DB list.

    Uses rapidfuzz.process.extractOne (token_sort_ratio) when available, which
    correctly handles name-part reordering (e.g. 'Abidi Mounir' vs 'Mounir ABIDI')
    and minor typos.  Falls back to substring matching when rapidfuzz is absent.

    Returns (matched_ids, name_to_match) where name_to_match maps
    fiche_name → (ens_id, display_name).
    """
    matched_ids: List[str] = []
    name_to_match: Dict[str, Tuple[str, str]] = {}

    if not enseignants:
        return matched_ids, name_to_match

    # Build lookup structures once
    # Each entry: (ens_id, display_name, normalised_full_name)
    ens_lookup: List[Tuple[str, str, str]] = [
        (e.id, f"{e.prenom} {e.nom}".strip(), _normalize(f"{e.prenom} {e.nom}"))
        for e in enseignants
    ]

    for fiche_name in fiche_names:
        fn_norm = _normalize(fiche_name)

        # ── rapidfuzz path ────────────────────────────────────────────────
        if _FUZZY_OK:
            choices = {idx: entry[2] for idx, entry in enumerate(ens_lookup)}
            best = _rprocess.extractOne(
                fn_norm,
                choices,
                scorer=_rfuzz.token_sort_ratio,
                score_cutoff=_FUZZY_THRESHOLD,
            )
            if best:
                _, score, idx = best
                eid, display, _ = ens_lookup[idx]
                matched_ids.append(eid)
                name_to_match[fiche_name] = (eid, display)
                logger.debug(f"  Fuzzy match '{fiche_name}' → '{display}' (score={score})")
            continue

        # ── substring fallback ────────────────────────────────────────────
        for eid, display, ens_norm in ens_lookup:
            ens_nom_norm = _normalize(eid.split("-")[-1]) if "-" in eid else ""
            if ens_norm in fn_norm or _normalize("".join(reversed(ens_norm.split()))) in fn_norm:
                matched_ids.append(eid)
                name_to_match[fiche_name] = (eid, display)
                break
            # Last-name only (>3 chars)
            last = ens_norm.split()[-1] if ens_norm.split() else ""
            if len(last) > 3 and last in fn_norm:
                matched_ids.append(eid)
                name_to_match[fiche_name] = (eid, display)
                break

    return matched_ids, name_to_match

def _match_enseignants_by_module(
    text: str,
    enseignants: List[EnseignantInfo],
) -> List[str]:
    """Match enseignants whose declared modules overlap with the fiche text."""
    norm_text = _normalize(text)
    matched = []
    for ens in enseignants:
        for mod in ens.modules:
            if len(mod) > 3 and _normalize(mod) in norm_text:
                matched.append(ens.id)
                break
    return matched

# ─────────────────────────────────────────────────────────────────────────────
# Core NLP analyzer – builds the competence tree from extracted data
# ─────────────────────────────────────────────────────────────────────────────

def _analyze_single_fiche(
    filename: str,
    text: str,
    enseignants: List[EnseignantInfo],
    departement: str = "gc",
    raw_tables: Optional[List] = None,
) -> Tuple[DomaineProposition, List[FicheEnseignantExtrait]]:
    """
    Full NLP pipeline for one fiche module:
      1. Extract metadata (NER — table-based primary, regex fallback)
      2. Extract Acquis d'Apprentissage (Bloom taxonomy)
      3. Extract Séances (content chunking)
      4. Build competence tree
      5. Match enseignants (fuzzy name + module matching)
    """
    logger.info(f"Analyzing fiche: {filename} ({len(text)} chars) [dept={departement}]")

    # ── Step 1: Metadata extraction ──────────────────────────────────────
    meta = _extract_metadata(text, raw_tables=raw_tables)
    logger.info(f"  Metadata: {meta}")

    domain_name = meta.get("unite_pedagogique")
    module_name = meta.get("nom_module")
    module_code = meta.get("code_module")

    if not domain_name:
        stem = Path(filename).stem
        clean = re.sub(r"^(fiche[_\s]?)?(ue[_\s]?|module[_\s]?|cours[_\s]?)?",
                       "", stem, flags=re.IGNORECASE)
        domain_name = re.sub(r"[_\-]+", " ", clean).strip().title() or "Domaine G\u00e9n\u00e9ral"

    if not module_name:
        module_name = domain_name
    if not module_code:
        module_code = _slug(module_name, 15)

    domain_code = _slug(domain_name, 10)

    # ── Step 2: Enseignant matching ──────────────────────────────────────
    fiche_ens_names = meta.get("enseignants_noms", [])
    roles_map = meta.get("enseignants_roles", {})
    if meta.get("responsable"):
        if meta["responsable"] not in fiche_ens_names:
            fiche_ens_names.append(meta["responsable"])
        roles_map.setdefault(meta["responsable"], "responsable")
    fiche_ens_names = list(dict.fromkeys(fiche_ens_names))  # dedupe preserving order

    matched_by_name, name_match_map = _match_enseignants_by_name(fiche_ens_names, enseignants)
    matched_by_module = _match_enseignants_by_module(text, enseignants)

    # Build extracted enseignant list – every name gets a real DB ID
    extracted_ens: List[FicheEnseignantExtrait] = []
    for name in fiche_ens_names:
        role = roles_map.get(name, "enseignant")
        match_info = name_match_map.get(name)
        if match_info:
            mid, mnom = match_info
        else:
            # Name not found in DB → left unmatched for manual linking in the UI
            mid, mnom = None, None
        extracted_ens.append(FicheEnseignantExtrait(
            fichier=filename,
            nom_complet=name,
            role=role,
            matched_id=mid,
            matched_nom=mnom,
        ))
    logger.info(f"  Extracted professor names: {[e.nom_complet for e in extracted_ens]}")

    # ── Add referential & module-matched teachers (by ID) to extractedEnseignants ──
    # Build a name-lookup map: id → EnseignantInfo from passed list + DB cache
    all_ens_by_id: Dict[str, "EnseignantInfo"] = {str(e.id): e for e in enseignants}
    try:
        all_ens_by_id.update(_fetch_all_enseignants_info())
    except Exception:
        pass  # DB unreachable – use only the passed list

    # Referential-based matching: match text → savoirs → enseignants
    ref_savoir_matches = _match_gc_savoir(text, departement=departement)
    matched_by_ref = _suggest_gc_enseignants(ref_savoir_matches[:10])
    all_matched = list(set(matched_by_name + matched_by_module + matched_by_ref))

    # Enrich extractedEnseignants: add any matched ID that is not yet represented
    already_extracted_ids = {ex.matched_id for ex in extracted_ens if ex.matched_id}
    for eid in all_matched:
        eid_str = str(eid)
        if eid_str in already_extracted_ids:
            continue
        ens_obj = all_ens_by_id.get(eid_str)
        if not ens_obj:
            # Cannot resolve a real name – skip rather than showing a raw code
            continue
        full_name = f"{ens_obj.prenom} {ens_obj.nom}".strip() or eid_str
        extracted_ens.append(FicheEnseignantExtrait(
            fichier=filename,
            nom_complet=full_name,
            role="enseignant",
            matched_id=eid_str,
            matched_nom=full_name,
        ))
        already_extracted_ids.add(eid_str)

    logger.info(f"  Matched enseignants: {len(all_matched)} "
                f"(by name: {len(matched_by_name)}, by module: {len(matched_by_module)}, "
                f"by ref [{departement}]: {len(matched_by_ref)})")
    if ref_savoir_matches:
        logger.info(f"  Referential savoir matches [{departement}]: {ref_savoir_matches[:5]}")

    # IDs of professors extracted directly from this fiche (last-resort fallback)
    extracted_ids: List[str] = [e.matched_id for e in extracted_ens if e.matched_id]

    # ── Step 3: Extract Acquis d'Apprentissage ───────────────────────────
    acquis = _extract_acquis_apprentissage(text)
    logger.info(f"  Acquis d'apprentissage found: {len(acquis)}")

    # ── Step 4: Extract Séances ──────────────────────────────────────────
    seances = _extract_seances(text)
    logger.info(f"  S\u00e9ances found: {len(seances)}")

    # ── Step 5: Build competence tree ────────────────────────────────────
    competences: List[CompetenceProposition] = []
    comp_idx = 0

    # ── Step 5a: Extract explicit sub-competence titles from the fiche ────
    subcomp_titles = _extract_subcompetences(text)
    if not subcomp_titles:
        llm_titles = _llm_extract_subcompetences(text, module_name)
        if llm_titles:
            subcomp_titles = [(i + 1, t) for i, t in enumerate(llm_titles)]
    if subcomp_titles:
        logger.info(f"  Sous-compétences explicites trouvées: {len(subcomp_titles)} — "
                     f"{[t for _, t in subcomp_titles[:5]]}")

    # === Compétence 1: from Acquis d'Apprentissage ===
    if acquis:
        comp_code = f"{module_code}-C{comp_idx + 1}"
        sous_comps: List[SousCompetenceProposition] = []
        sc_idx = 0

        if subcomp_titles:
            # ── Dynamic sub-competence creation from explicit titles ──────
            text_lines = text.splitlines()
            for idx_sc, (line_no, title) in enumerate(subcomp_titles):
                sc_code = f"{comp_code}-SC{sc_idx + 1}"
                # Determine text range for this sub-competence
                start = line_no - 1  # 0-based
                end = (subcomp_titles[idx_sc + 1][0] - 1) if (idx_sc + 1 < len(subcomp_titles)) else len(text_lines)
                sc_block_text = "\n".join(text_lines[start:end])

                # Find AAs whose text appears in this block
                aa_block = [aa for aa in acquis if aa["text"] in sc_block_text]
                if not aa_block:
                    aa_block = acquis  # fallback: all AAs

                savoirs = []
                for j, aa in enumerate(aa_block):
                    sav_code = f"{sc_code}-S{j + 1}"
                    gc_codes = _match_gc_savoir(aa["text"], departement=departement)
                    savoir_ens = list(
                        set(matched_by_name + matched_by_module + _suggest_gc_enseignants(gc_codes))
                    ) if gc_codes else list(set(matched_by_name + matched_by_module))
                    savoirs.append(SavoirProposition(
                        tmpId=str(uuid.uuid4()),
                        code=sav_code,
                        nom=aa["text"][:120],
                        description=aa["text"][:200],
                        type=_detect_type(aa["text"], departement),
                        niveau=_gc_ref_niveau(gc_codes, departement=departement) or _bloom_to_niveau(aa["bloom_level"]),
                        enseignantsSuggeres=savoir_ens or all_matched or extracted_ids,
                        refCodes=gc_codes,
                    ))

                sc_gc = list({c for s in savoirs for c in s.refCodes})
                sous_comps.append(SousCompetenceProposition(
                    tmpId=str(uuid.uuid4()),
                    code=sc_code,
                    nom=title,
                    description=None,
                    refCodes=sc_gc,
                    savoirs=savoirs,
                ))
                sc_idx += 1
        else:
            # ── Fallback: group AAs by Bloom level (existing behaviour) ──
            groups: Dict[str, List[Dict]] = {
                "Connaissances fondamentales": [],
                "Compétences appliquées": [],
                "Compétences avancées": [],
            }
            for aa in acquis:
                bl = aa["bloom_level"]
                if bl <= 2:
                    groups["Connaissances fondamentales"].append(aa)
                elif bl <= 4:
                    groups["Compétences appliquées"].append(aa)
                else:
                    groups["Compétences avancées"].append(aa)

            for sc_name, aa_list in groups.items():
                if not aa_list:
                    continue
                sc_code = f"{comp_code}-SC{sc_idx + 1}"
                savoirs = []
                for j, aa in enumerate(aa_list):
                    sav_code = f"{sc_code}-S{j + 1}"
                    gc_codes = _match_gc_savoir(aa["text"], departement=departement)
                    savoir_ens = list(
                        set(matched_by_name + matched_by_module + _suggest_gc_enseignants(gc_codes))
                    ) if gc_codes else list(set(matched_by_name + matched_by_module))
                    savoirs.append(SavoirProposition(
                        tmpId=str(uuid.uuid4()),
                        code=sav_code,
                        nom=aa["text"][:120],
                        description=aa["text"][:200],
                        type=_detect_type(aa["text"], departement),
                        niveau=_gc_ref_niveau(gc_codes, departement=departement) or _bloom_to_niveau(aa["bloom_level"]),
                        enseignantsSuggeres=savoir_ens or all_matched or extracted_ids,
                        refCodes=gc_codes,
                    ))
                sc_gc = list({c for s in savoirs for c in s.refCodes})
                sous_comps.append(SousCompetenceProposition(
                    tmpId=str(uuid.uuid4()),
                    code=sc_code,
                    nom=sc_name,
                    description=None,
                    refCodes=sc_gc,
                    savoirs=savoirs,
                ))
                sc_idx += 1

        # Aggregate GC codes from sous-compétences and classify at competence level
        comp_gc = list({c for sc2 in sous_comps for c in sc2.refCodes})
        comp_gc_domaine = _match_gc_competence(" ".join(comp_gc), departement=departement) if comp_gc else None
        competences.append(CompetenceProposition(
            tmpId=str(uuid.uuid4()),
            code=comp_code,
            nom=f"Acquis d'apprentissage \u2013 {module_name}",
            description=meta.get("objectif"),
            ordre=comp_idx + 1,
            refCodes=comp_gc,
            refDomaine=comp_gc_domaine,
            sousCompetences=sous_comps,
        ))
        comp_idx += 1

    # === Compétence(s) from Séances (contenu détaillé) ===
    if seances:
        comp_code = f"{module_code}-C{comp_idx + 1}"
        sous_comps_seances: List[SousCompetenceProposition] = []

        for sc_idx, seance in enumerate(seances):
            sc_code = f"{comp_code}-SC{sc_idx + 1}"
            items = seance["items"]
            if not items:
                items = [seance["titre"]]

            savoirs = []
            for j, item in enumerate(items):
                sav_code = f"{sc_code}-S{j + 1}"
                bloom = _detect_bloom_level(item)
                item_type = _detect_type(item + " " + (seance.get("type_apprentissage") or ""), departement)
                gc_codes = _match_gc_savoir(item, departement=departement)
                # Per-savoir teacher suggestions: name/module matches + ref-code-specific matches
                savoir_ens = list(
                    set(matched_by_name + matched_by_module + _suggest_gc_enseignants(gc_codes))
                ) if gc_codes else list(set(matched_by_name + matched_by_module))
                savoirs.append(SavoirProposition(
                    tmpId=str(uuid.uuid4()),
                    code=sav_code,
                    nom=item[:120],
                    description=None,
                    type=item_type,
                    niveau=_gc_ref_niveau(gc_codes, departement=departement) or _bloom_to_niveau(bloom),
                    enseignantsSuggeres=savoir_ens or all_matched or extracted_ids,
                    refCodes=gc_codes,
                ))

            sc_titre = f"S\u00e9ance {seance['numero']} : {seance['titre']}"
            if seance.get("duree"):
                sc_titre += f" ({seance['duree']})"

            sc_gc = list({c for s in savoirs for c in s.refCodes})
            sous_comps_seances.append(SousCompetenceProposition(
                tmpId=str(uuid.uuid4()),
                code=sc_code,
                nom=sc_titre[:100],
                description=seance.get("type_apprentissage"),
                refCodes=sc_gc,
                savoirs=savoirs,
            ))

        seance_gc = list({c for sc2 in sous_comps_seances for c in sc2.refCodes})
        seance_gc_domaine = _match_gc_competence(" ".join(seance_gc), departement=departement) if seance_gc else None
        competences.append(CompetenceProposition(
            tmpId=str(uuid.uuid4()),
            code=comp_code,
            nom=f"Contenu d\u00e9taill\u00e9 \u2013 {module_name}",
            description=None,
            ordre=comp_idx + 1,
            refCodes=seance_gc,
            refDomaine=seance_gc_domaine,
            sousCompetences=sous_comps_seances,
        ))
        comp_idx += 1

    # === Fallback: generic extraction if no AA and no séances found ===
    if not competences:
        competences = _fallback_extraction(text, module_code, module_name,
                                           all_matched or extracted_ids,
                                           departement=departement)

    # Aggregate all referential codes from the full competence tree
    all_ref_fiche = list({
        c
        for comp in competences
        for sc2 in comp.sousCompetences
        for s in sc2.savoirs
        for c in s.refCodes
    })
    domaine_ref_domaine = _match_gc_competence(text, departement=departement)  # classify whole fiche
    if domaine_ref_domaine:
        logger.info(f"  Domain match [{departement}]: {domaine_ref_domaine} ({len(all_ref_fiche)} ref codes)")

    domaine = DomaineProposition(
        tmpId=str(uuid.uuid4()),
        code=domain_code,
        nom=domain_name,
        description=f"Domaine extrait de : {filename}" + (
            f" | Pr\u00e9requis: {meta['prerequis']}" if meta.get("prerequis") else ""
        ),
        refCodes=all_ref_fiche,
        refDomaine=domaine_ref_domaine,
        competences=competences,
    )
    return domaine, extracted_ens


def _fallback_extraction(
    text: str,
    module_code: str,
    module_name: str,
    matched_ens: List[str],
    departement: str = "gc",
) -> List[CompetenceProposition]:
    """
    Fallback when structured sections (AA, Séances) are not found.

    Strategy:
    1. **LLM** (primary) — extracts action-verb competence phrases from any text.
    2. **Regex + Bloom filter** (fallback) — bullet extraction with verb validation.
    """
    items: List[str] = []

    # ── 0. LLM extraction (primary) ──────────────────────────────────────
    if _LLM_OK:
        items = _llm_fallback_items(text, module_name)

    # ── 1. Regex + Bloom filter (fallback when LLM unavailable/empty) ─────
    if not items:
        for pattern in [_RE_CHECKMARK, _RE_BULLET, _RE_NUMBERED]:
            items.extend(pattern.findall(text))
        items = [it.strip() for it in items if len(it.strip()) > 5]

        # ── Bloom-verb filter: reject administrative/descriptive sentences ─────
        # Forbidden openers that are never competences
        _FALLBACK_NOISE = re.compile(
            r"^(ce\s+(module|cours|ue)|l['\'']objectif|les\s+(étudiants?|apprenants?)|"
            r"ce\s+cours\s+(est|comporte|comprend)|il\s+(est|s\'agit|faut)|"
            r"avoir|être|\d+\s*h)",
            re.IGNORECASE,
        )

        if not items:
            sentences = re.split(r"[.\n]+", text)
            items = [
                s.strip() for s in sentences
                if 10 < len(s.strip()) < 200
                and not re.match(r"^(Code|Mode|Evaluation|R\u00e9f\u00e9rence)", s.strip())
                and not _FALLBACK_NOISE.match(s.strip())
                # Require at least a level-2 Bloom verb (action word) to be a competence
                and _detect_bloom_level(s.strip()) >= 2
            ]
            items = items[:20]

    if not items:
        items = [module_name]

    comp_code = f"{module_code}-C1"
    sc_code = f"{comp_code}-SC1"

    savoirs = []
    for j, item in enumerate(items[:25]):
        sav_code = f"{sc_code}-S{j + 1}"
        bloom = _detect_bloom_level(item)
        gc_codes = _match_gc_savoir(item, departement=departement)
        savoir_ens = list(set(matched_ens + _suggest_gc_enseignants(gc_codes))) if gc_codes else matched_ens
        savoirs.append(SavoirProposition(
            tmpId=str(uuid.uuid4()),
            code=sav_code,
            nom=item[:120],
            description=None,
            type=_detect_type(item, departement),
            niveau=_gc_ref_niveau(gc_codes, departement=departement) or _bloom_to_niveau(bloom),
            enseignantsSuggeres=savoir_ens,
            refCodes=gc_codes,
        ))

    sc_gc = list({c for s in savoirs for c in s.refCodes})
    sc = SousCompetenceProposition(
        tmpId=str(uuid.uuid4()),
        code=sc_code,
        nom="Contenus extraits",
        refCodes=sc_gc,
        savoirs=savoirs,
    )

    comp_gc = list({c for s in savoirs for c in s.refCodes})
    comp_gc_domaine = _match_gc_competence(" ".join(comp_gc), departement=departement) if comp_gc else None
    return [CompetenceProposition(
        tmpId=str(uuid.uuid4()),
        code=comp_code,
        nom=module_name,
        ordre=1,
        refCodes=comp_gc,
        refDomaine=comp_gc_domaine,
        sousCompetences=[sc],
    )]


def analyze_files(
    filenames: List[str],
    file_contents: List[bytes],
    enseignants: List[EnseignantInfo],
    departement: str = "gc",
) -> RiceAnalysisResult:
    """Main analysis entry point – processes multiple fiche files.

    ``departement`` selects which referential is used for code matching:
    * ``"gc"``      – Génie Civil (built-in fallback + DB overrides)
    * ``"info"``    – Informatique (from DB ``ref_savoirs`` table)
    * ``"ge"``      – Génie Électrique (from DB)
    * ``"meca"``    – Génie Mécanique (from DB)
    * ``"telecom"`` – Télécommunications (from DB)
    * any other code – uses DB ``ref_savoirs`` rows filtered by that code
    """
    domaines: List[DomaineProposition] = []
    all_extracted_ens: List[FicheEnseignantExtrait] = []
    seen_codes: Dict[str, int] = {}

    for filename, data in zip(filenames, file_contents):
        text, raw_tables = _extract_text(filename, data)
        domaine, extracted_ens = _analyze_single_fiche(
            filename, text, enseignants,
            departement=departement,
            raw_tables=raw_tables,
        )

        # Deduplicate domain codes across files
        if domaine.code in seen_codes:
            seen_codes[domaine.code] += 1
            domaine.code = f"{domaine.code}{seen_codes[domaine.code]}"
        else:
            seen_codes[domaine.code] = 1

        domaines.append(domaine)
        all_extracted_ens.extend(extracted_ens)

    # Build summary stats
    total_comp = sum(len(d.competences) for d in domaines)
    total_sc   = sum(len(c.sousCompetences)
                     for d in domaines for c in d.competences)
    total_sav  = sum(len(sc.savoirs)
                     for d in domaines for c in d.competences
                     for sc in c.sousCompetences)
    assigned_ens = {
        eid
        for d in domaines for c in d.competences
        for sc in c.sousCompetences for s in sc.savoirs
        for eid in s.enseignantsSuggeres
    }

    all_ref_codes_covered = sorted({
        c
        for d in domaines
        for comp in d.competences
        for sc in comp.sousCompetences
        for s in sc.savoirs
        for c in s.refCodes
    })
    stats = {
        "departement": departement,
        "totalDomaines": len(domaines),
        "totalCompetences": total_comp,
        "totalSousCompetences": total_sc,
        "totalSavoirs": total_sav,
        "enseignantsCovered": len(assigned_ens),
        "tauxCouverture": round(len(assigned_ens) / max(len(enseignants), 1) * 100, 1),
        "refCodesCovered": all_ref_codes_covered,
        "refCodesCoveredCount": len(all_ref_codes_covered),
    }

    # Build list of detailed EnseignantInfo for any assigned ID found in DB
    all_db_infos = _fetch_all_enseignants_info()
    found_ens_list = []
    for eid in assigned_ens:
        if eid in all_db_infos:
            found_ens_list.append(all_db_infos[eid])

    return RiceAnalysisResult(
        propositions=domaines,
        stats=stats,
        extractedEnseignants=all_extracted_ens,
        foundEnseignants=found_ens_list,
    )


# ─────────────────────────────────────────────────────────────────────────────
# FastAPI endpoint
# ─────────────────────────────────────────────────────────────────────────────

import json as _json

# ─────────────────────────────────────────────────────────────────────────────
# Département-specific Referential Knowledge Base
# _GC_FALLBACK_REF is the built-in referential for ESPRIT Génie Civil.
# All other departments (info, ge, meca, telecom, …) load their referential
# exclusively from the `ref_savoirs` DB table (filtered by the `departement`
# column when present).
# To add a new department: INSERT rows into ref_savoirs with
# departement='<code>', code, nom, keywords.
# ─────────────────────────────────────────────────────────────────────────────

_GC_FALLBACK_REF = {
    # ── Domaines ────────────────────────────────────────────────────────────
    "domaines": {
        "GC-RDI":  "RDI – Recherche, Développement et Innovation",
        "GC-PERS": "Personnel et Relationnel",
        "GC-COM":  "Communication et Culture",
        "GC-PED":  "Pédagogie",
        "GC-TECH": "Technique / Métier Génie Civil",
    },

    # ── Compétences (6 compétences techniques) ──────────────────────────────
    "competences": {
        "GC-TECH-S": {
            "nom": "Compétences dans le domaine des sols (S)",
            "keywords": [
                "sol", "geologie", "geotechnique", "coupe geologique",
                "fondation", "soutenement", "pente", "sismique",
                "instabilite hydraulique", "renard", "boulance",
            ],
        },
        "GC-TECH-C": {
            "nom": "Compétences dans le domaine de la construction (C)",
            "keywords": [
                "beton arme", "ouvrage art", "pont", "viaduc",
                "infrastructure routiere", "route", "chaussee",
                "rehabilitation", "mode constructif", "prefabrique",
            ],
        },
        "GC-TECH-P": {
            "nom": "Compétences en physique du bâtiment (P)",
            "keywords": [
                "thermique", "acoustique", "aeraulique", "isolation",
                "physique batiment", "performance energetique",
                "equipement technique", "cvc", "ventilation",
            ],
        },
        "GC-TECH-E": {
            "nom": "Compétences dans le domaine de l'eau (E)",
            "keywords": [
                "hydraulique", "hydrologie", "bassin versant", "debit",
                "crue", "assainissement", "reseau eau", "diagnostic eau",
            ],
        },
        "GC-TECH-U": {
            "nom": "Compétences en urbanisme (U)",
            "keywords": [
                "urbanisme", "amenagement urbain", "diagnostic urbain",
                "situation urbaine", "ville", "territoire", "paysage",
            ],
        },
        "GC-TECH-T": {
            "nom": "Compétences transversales en génie civil (T)",
            "keywords": [
                "pluridisciplinaire", "organisation chantier", "securite chantier",
                "construction durable", "developpement durable",
                "maintenance ouvrage", "assurance qualite", "plan qualite",
            ],
        },
    },

    # ── Savoirs : code → [mots-clés de matching] ───────────────────────────
    # Chaque entrée reflète exactement l'intitulé officiel du référentiel GC
    "savoirs": {
        # ── S – Sols ──────────────────────────────────────────────────────
        "S1a": ["coupe geologique", "effectuer coupe", "coupe lithologique"],
        "S1b": ["interpreter coupe geologique", "coupe geologique interpreter",
                "interpretation geologique"],
        "S1c": ["teledetection", "carte geologique", "interpreter carte",
                "resultat teledetection"],
        "S1d": ["horizon geologique", "identifier horizon", "identification couche"],
        "S2a": ["essai geotechnique laboratoire", "essai classification",
                "comportement sol laboratoire", "realiser essai geotechnique"],
        "S2b": ["interpreter essai geotechnique", "resultat essai geotechnique",
                "interpretation laboratoire"],
        "S3":  ["rupture pente", "stabilite pente", "glissement terrain",
                "risque pente", "sollicitation pente"],
        "S4":  ["instabilite hydraulique sol", "risque hydraulique sol",
                "renard", "boulance", "soulevement hydraulique"],
        "S5":  ["risque geotechnique sismique", "sollicitation sismique",
                "risque sismique geotechnique", "seisme sol"],
        "S6a": ["concevoir fondation", "concevoir soutenement",
                "systeme fondation conception", "paroi soutenement conception"],
        "S6b": ["dimensionner fondation", "dimensionner soutenement",
                "calcul fondation", "pieu calcul", "semelle dimensionner"],
        "S6c": ["controler fondation", "controler soutenement",
                "verification fondation", "reception fondation"],
        # ── C – Construction ──────────────────────────────────────────────
        "C1a": ["concevoir structure beton", "beton arme concevoir",
                "structure batiment conception", "conception batiment beton"],
        "C1b": ["dimensionner beton arme", "calcul structure beton",
                "bael", "eurocode 2", "dimensionnement beton"],
        "C1c": ["controler beton arme", "verification beton arme",
                "conformite structurale", "inspection beton"],
        "C2a": ["concevoir ouvrage art", "pont conception", "viaduc conception",
                "ouvrage art concevoir"],
        "C2b": ["dimensionner ouvrage art", "calcul pont", "dimensionner pont",
                "calcul viaduc"],
        "C2c": ["controler ouvrage art", "verification pont", "reception pont",
                "inspection ouvrage art"],
        "C3a": ["concevoir infrastructure routiere", "trace routier",
                "route conception", "conception route"],
        "C3b": ["dimensionner infrastructure routiere", "chaussee dimensionner",
                "terrassement", "dimensionner route"],
        "C3c": ["controler infrastructure routiere", "chantier routier",
                "reception chaussee", "supervision route"],
        "C4":  ["gestion projet infrastructure", "management projet infrastructure",
                "pilotage projet", "chef projet infrastructure"],
        "C5":  ["etude impact infrastructure", "impact environnement infrastructure",
                "etude impact"],
        "C6":  ["mode constructif", "methode construction", "prefabrique",
                "coffrage", "choisir mode constructif"],
        "C7":  ["etat sante structurel", "sante structurel", "diagnostic structure",
                "pathologie batiment", "actions necessaires structure"],
        "C8":  ["rehabilitation ouvrage art", "actions rehabilitation",
                "reparation ouvrage art", "refection ouvrage"],
        # ── P – Physique du bâtiment ──────────────────────────────────────
        "P1a": ["concevoir solution thermique", "concevoir solution acoustique",
                "aeraulique conception", "physique batiment conception"],
        "P1b": ["dimensionner solution thermique", "dimensionner solution acoustique",
                "calcul thermique batiment", "aeraulique dimensionner"],
        "P1c": ["controler solution thermique", "controler solution acoustique",
                "performance thermique verification", "verification acoustique"],
        "P2":  ["diagnostic thermique batiment", "performance thermique evaluation",
                "acoustique evaluation", "bilan thermique batiment", "etat sante thermique"],
        "P3":  ["dimensionner equipement technique", "choisir equipement technique",
                "cvc", "plomberie", "ventilation dimensionner"],
        # ── E – Eau ───────────────────────────────────────────────────────
        "E1a": ["concevoir reseau hydraulique", "concevoir ouvrage hydraulique",
                "hydraulique urbain conception", "hydrologie reseau conception"],
        "E1b": ["dimensionner reseau hydraulique", "dimensionner ouvrage hydraulique",
                "calcul hydraulique", "dimensionner reseau assainissement"],
        "E2":  ["diagnostic hydrologie quantitative", "gestion reseau hydraulique",
                "diagnostic reseau hydraulique", "hydrologie gestion"],
        "E3":  ["diagnostic environnemental eau", "systeme gestion eaux",
                "traitement eaux", "gestion dechets eau", "assainissement diagnostic"],
        # ── U – Urbanisme ─────────────────────────────────────────────────
        "U1":  ["analyser situation urbaine", "analyse urbaine", "diagnostic urbain",
                "situation technique urbaine", "echelle urbaine"],
        "U2":  ["realiser diagnostic urbain", "etude urbaine", "diagnostic urbain"],
        "U3a": ["concevoir amenagement urbain", "projet amenagement urbain",
                "plan amenagement", "design urbain"],
        "U3b": ["conduire projet amenagement urbain", "piloter amenagement",
                "mise en oeuvre amenagement", "gestion projet urbain"],
        # ── T – Transversales ─────────────────────────────────────────────
        "T1":  ["conception pluridisciplinaire", "pluridisciplinaire batiment",
                "interaction architecture sol", "integration disciplines"],
        "T2":  ["organisation chantier", "procedes construction",
                "securite chantier", "maitrise delais", "chef chantier"],
        "T3":  ["construction durable", "amenagement durable",
                "developpement durable construction", "hqe", "eco construction"],
        "T4":  ["gestion ouvrage existant", "maintenance ouvrage",
                "evaluer ouvrage", "maintenir ouvrage", "patrimoine ouvrage"],
        "T5":  ["assurance qualite", "plan qualite", "aqp",
                "management qualite", "normes qualite"],
    },

    # ── Niveaux officiels par savoir (N1=débutant … N5=expert) ────────────
    # Source : tableau "Affectation des savoirs par compétence et par niveaux"
    "niveaux": {
        # Sols
        "S2a": "N1_DEBUTANT",
        "S1a": "N2_ELEMENTAIRE",
        "S2b": "N2_ELEMENTAIRE",
        "S1b": "N3_INTERMEDIAIRE",
        "S1c": "N3_INTERMEDIAIRE",
        "S6b": "N3_INTERMEDIAIRE",
        "S1d": "N4_AVANCE",
        "S6a": "N4_AVANCE",
        "S3":  "N5_EXPERT",
        "S4":  "N5_EXPERT",
        "S5":  "N5_EXPERT",
        "S6c": "N5_EXPERT",
        # Construction
        "C1b": "N3_INTERMEDIAIRE",
        "C2b": "N3_INTERMEDIAIRE",
        "C3b": "N3_INTERMEDIAIRE",
        "C4":  "N4_AVANCE",
        "C1c": "N4_AVANCE",
        "C1a": "N4_AVANCE",
        "C2c": "N4_AVANCE",
        "C3a": "N4_AVANCE",
        "C3c": "N4_AVANCE",
        "C2a": "N4_AVANCE",
        "C5":  "N5_EXPERT",
        "C6":  "N5_EXPERT",
        "C7":  "N5_EXPERT",
        "C8":  "N5_EXPERT",
        # Physique du bâtiment
        "P1b": "N3_INTERMEDIAIRE",
        "P2":  "N3_INTERMEDIAIRE",
        "P3":  "N4_AVANCE",
        "P1c": "N4_AVANCE",
        "P1a": "N5_EXPERT",
        # Eau
        "E1b": "N4_AVANCE",
        "E1a": "N5_EXPERT",
        "E2":  "N5_EXPERT",
        "E3":  "N5_EXPERT",
        # Urbanisme
        "U1":  "N3_INTERMEDIAIRE",
        "U2":  "N3_INTERMEDIAIRE",
        "U3a": "N5_EXPERT",
        "U3b": "N5_EXPERT",
        # Transversales
        "T2":  "N4_AVANCE",
        "T1":  "N5_EXPERT",
        "T3":  "N5_EXPERT",
        "T4":  "N5_EXPERT",
        "T5":  "N5_EXPERT",
    },
}


def _match_gc_savoir(text: str, departement: str = "gc") -> List[str]:
    """
    Match text against the department referential to find matching savoir codes.

    Strategy:
      1. **Semantic matching** (primary) — sentence-transformer embeddings capture
         contextual meaning and fill vocabulary gaps.
      2. **Keyword matching** (high-confidence override) — exact keyword hits are
         authoritative; they are placed first when present, with any unique semantic
         results appended after them.

    Returns a deduplicated list of savoir codes sorted by relevance.
    """
    ref = _get_effective_referential(departement)
    norm = _normalize(text)

    # --- Primary: semantic similarity ------------------------------------------
    semantic_codes: List[str] = []
    if _SEMANTIC_OK:
        semantic_codes = _match_gc_savoir_semantic(text, departement=departement)
        if semantic_codes:
            logger.debug(f"Semantic [{departement}]: '{text[:60]}' → {semantic_codes}")

    # --- Boost: keyword exact-match (higher confidence) ------------------------
    matches = []
    for code, keywords in ref["savoirs"].items():
        score = sum(1 for kw in keywords if kw in norm)
        if score > 0:
            matches.append((code, score))
    matches.sort(key=lambda x: x[1], reverse=True)
    keyword_codes = [code for code, _ in matches]

    if keyword_codes:
        # Keywords are high-precision → put them first, then semantic extras
        seen = set(keyword_codes)
        return keyword_codes + [c for c in semantic_codes if c not in seen]

    # No keyword hits: rely entirely on semantic similarity
    return semantic_codes


def _gc_ref_niveau(gc_codes: List[str], departement: str = "gc") -> Optional[str]:
    """
    Return the official niveau for the best-matched savoir code.
    Prioritises the highest-scoring (first) code in the list.
    Returns None if no code is found in the niveaux table.
    """
    niveaux_map = _get_effective_referential(departement).get("niveaux", {})
    for code in gc_codes:
        if code in niveaux_map:
            return niveaux_map[code]
    return None


def _match_gc_competence(text: str, departement: str = "gc") -> Optional[str]:
    """
    Match text against the department referential to find the best-matching competence code.
    Uses the active (DB-aware) referential for the given department.
    """
    ref = _get_effective_referential(departement)
    norm = _normalize(text)
    best_code = None
    best_score = 0
    for code, info in ref["competences"].items():
        score = sum(1 for kw in info["keywords"] if kw in norm)
        if score > best_score:
            best_score = score
            best_code = code
    return best_code if best_score > 0 else None


def _suggest_gc_enseignants(savoir_codes: List[str]) -> List[str]:
    """
    Given a list of savoir codes, find which GC enseignants are mapped to them.
    Fetches the mapping dynamically from the database.
    """
    affectations = _fetch_enseignant_affectations()
    suggested = set()
    for ens_id, codes in affectations.items():
        if any(_codes_match(ec, sc) for ec in codes for sc in savoir_codes):
            suggested.add(ens_id)
    return list(suggested)


# ─────────────────────────────────────────────────────────────────────────────
# Point 3 – DB-backed Referential (multi-department, merges DB + per-dept fallback)
# ─────────────────────────────────────────────────────────────────────────────

# Keyed by lowercase department code (e.g. "gc", "info", "ge")
_REF_DB_CACHE = _ThreadSafeCache()
_REF_DB_TTL: float = 600.0  # 10 minutes

# Empty baseline returned for unknown departments with no DB entries
_EMPTY_REFERENTIAL: Dict = {"domaines": {}, "competences": {}, "savoirs": {}, "niveaux": {}}

# Minimal generic referential for non-GC departments that have no DB data yet.
# Provides the structural skeleton so the competence tree is never fully empty.
_GENERIC_FALLBACK_REF: Dict = {
    "domaines":    {"GEN": "Compétences Générales"},
    "competences": {},
    "savoirs":     {},
    "niveaux": {
        "N1_DEBUTANT":      {"label": "Débutant",      "score": 1},
        "N2_ELEMENTAIRE":   {"label": "Élémentaire",   "score": 2},
        "N3_INTERMEDIAIRE": {"label": "Intermédiaire", "score": 3},
        "N4_AVANCE":        {"label": "Avancé",        "score": 4},
        "N5_EXPERT":        {"label": "Expert",        "score": 5},
    },
}


def _load_ref_from_db(departement: str = "gc") -> Optional[Dict]:
    """Load the full department referential from DB: savoirs + competences + domaines.

    Tables queried (all optional – graceful degradation when absent):
    * ``ref_savoirs``    – savoir keywords, filtered by ``departement`` column when present
    * ``ref_competences``– competence keywords per department (created by migrate script)
    * ``ref_domaines``   – domain names per department (created by migrate script)

    Behaviour:
    * Génie Civil (``departement='gc'``): DB rows are merged **on top of** the
      built-in ``_GC_FALLBACK_REF`` so that the in-memory dictionary acts as a
      safety net when the DB table is empty or absent.
    * Any other department: only DB rows are used.  Populate the tables via
      ``seed_autres_departements.py`` (or the migration SQL script).

    Returns the merged referential dict, or ``None`` on DB error / table absent.
    """
    global _SEMANTIC_CORPUS_BUILT
    dept_key = departement.lower().strip()
    cached = _REF_DB_CACHE.get(dept_key, ttl=_REF_DB_TTL)
    if cached is not None:
        return cached
    try:
        conn = _get_db_connection()
        cur = conn.cursor()

        # ── 1. Check ref_savoirs table existence ──────────────────────────
        cur.execute("""
            SELECT EXISTS (
                SELECT 1 FROM information_schema.tables
                WHERE table_schema = 'public' AND table_name = 'ref_savoirs'
            )
        """)
        if not cur.fetchone()[0]:
            cur.close(); _put_db_connection(conn)
            return None

        # ── 2. Load savoirs (filter by departement column when present) ────
        cur.execute("""
            SELECT column_name FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = 'ref_savoirs'
              AND column_name = 'departement'
        """)
        has_dept_col = cur.fetchone() is not None
        if has_dept_col:
            cur.execute("SELECT code, nom, keywords FROM ref_savoirs WHERE departement = %s",
                        (dept_key,))
        else:
            cur.execute("SELECT code, nom, keywords FROM ref_savoirs")
        override: Dict[str, List] = {}
        for code, nom, keywords in cur.fetchall():
            if isinstance(keywords, list):
                kws = list(keywords)
            elif isinstance(keywords, str):
                kws = [k.strip().lower() for k in keywords.split(',') if k.strip()]
            else:
                kws = []
            if nom:
                kws_norm = [_normalize(k) for k in kws]
                nom_norm = _normalize(nom)
                if nom_norm not in kws_norm:
                    kws.append(nom.lower())
            override[code] = kws

        # ── 3. Load competences from ref_competences (if table exists) ─────
        db_competences: Dict[str, Any] = {}
        try:
            cur.execute("""
                SELECT EXISTS (
                    SELECT 1 FROM information_schema.tables
                    WHERE table_schema = 'public' AND table_name = 'ref_competences'
                )
            """)
            if cur.fetchone()[0]:
                cur.execute("""
                    SELECT code, nom, keywords
                    FROM ref_competences
                    WHERE departement = %s
                """, (dept_key,))
                for comp_code, comp_nom, comp_kws in cur.fetchall():
                    kws = comp_kws if isinstance(comp_kws, list) else (
                        [k.strip().lower() for k in (comp_kws or "").split(",") if k.strip()]
                    )
                    db_competences[comp_code] = {
                        "nom": comp_nom or comp_code,
                        "keywords": kws,
                    }
        except Exception as comp_exc:
            logger.debug(f"Cannot load ref_competences for [{dept_key}]: {comp_exc}")

        # ── 4. Load domaines from ref_domaines (if table exists) ───────────
        db_domaines: Dict[str, str] = {}
        try:
            cur.execute("""
                SELECT EXISTS (
                    SELECT 1 FROM information_schema.tables
                    WHERE table_schema = 'public' AND table_name = 'ref_domaines'
                )
            """)
            if cur.fetchone()[0]:
                cur.execute("""
                    SELECT code, nom FROM ref_domaines WHERE departement = %s
                """, (dept_key,))
                for dom_code, dom_nom in cur.fetchall():
                    db_domaines[dom_code] = dom_nom or dom_code
        except Exception as dom_exc:
            logger.debug(f"Cannot load ref_domaines for [{dept_key}]: {dom_exc}")

        cur.close(); _put_db_connection(conn)

        if override or db_competences or db_domaines:
            # For GC: merge DB data on top of in-memory fallback (safety net)
            if dept_key in ("gc", "genie_civil", "genie-civil"):
                base = _GC_FALLBACK_REF
                merged: Dict = {
                    **base,
                    "savoirs":     {**base["savoirs"],    **override},
                    "competences": {**base["competences"], **db_competences} if db_competences else base["competences"],
                    "domaines":    {**base["domaines"],    **db_domaines}    if db_domaines    else base["domaines"],
                }
            else:
                # Non-GC departments: DB is the sole source of truth
                merged = {
                    **_EMPTY_REFERENTIAL,
                    "savoirs":     override,
                    "competences": db_competences,
                    "domaines":    db_domaines,
                    "niveaux":     {},
                }
                # Also try loading niveaux from ref_savoirs.niveau column
                try:
                    conn2 = _get_db_connection()
                    cur2 = conn2.cursor()
                    cur2.execute("""
                        SELECT column_name FROM information_schema.columns
                        WHERE table_schema = 'public' AND table_name = 'ref_savoirs'
                          AND column_name = 'niveau'
                    """)
                    if cur2.fetchone():
                        cur2.execute("""
                            SELECT code, niveau FROM ref_savoirs
                            WHERE departement = %s AND niveau IS NOT NULL
                        """, (dept_key,))
                        for sav_code, sav_niveau in cur2.fetchall():
                            merged["niveaux"][sav_code] = sav_niveau
                    cur2.close(); _put_db_connection(conn2)
                except Exception:
                    pass

            _REF_DB_CACHE.set(dept_key, merged)
            logger.info(
                f"Referential loaded from DB [{dept_key}]: "
                f"{len(override)} savoirs, {len(db_competences)} compétences, "
                f"{len(db_domaines)} domaines"
            )
            _SEMANTIC_CORPUS_BUILT = False  # force re-encoding on dept change
            return merged
    except Exception as exc:
        logger.debug(f"Cannot load referential from DB for [{dept_key}] (ok if absent): {exc}")
    return None


# ── JSON-file-based generic referentials (fallback when DB is absent) ──────────
import json as _json_local

_GENERIC_REF_DIR = Path(__file__).parent / "refs"
_GENERIC_REF_MAPPING_PATH = _GENERIC_REF_DIR / "generic_ref.json"


def _load_generic_ref(departement: str) -> Dict:
    """Load a department referential from the JSON files in ``refs/``.

    Falls back to ``_GENERIC_FALLBACK_REF`` if the file is missing or corrupt.
    """
    try:
        mapping = _json_local.loads(_GENERIC_REF_MAPPING_PATH.read_text(encoding="utf-8"))
        rel_path = mapping.get(departement.lower().strip())
        if not rel_path:
            logger.info(f"No generic ref mapping entry for '{departement}'")
            return _GENERIC_FALLBACK_REF
        ref_file = _GENERIC_REF_DIR / Path(rel_path).name
        if not ref_file.is_file():
            # Also try the relative path as-is from the project root
            ref_file = Path(__file__).parent / rel_path
        if not ref_file.is_file():
            logger.info(f"Generic ref file not found for '{departement}': {ref_file}")
            return _GENERIC_FALLBACK_REF
        data = _json_local.loads(ref_file.read_text(encoding="utf-8"))
        # Validate minimal structure
        for key in ("domaines", "competences", "savoirs", "niveaux"):
            if key not in data:
                data[key] = {}
        logger.info(f"Generic ref loaded from JSON for '{departement}': "
                    f"{len(data.get('savoirs', {}))} savoirs")
        return data
    except Exception as exc:
        logger.warning(f"Impossible de charger le référentiel générique pour '{departement}': {exc}")
        return _GENERIC_FALLBACK_REF


def _get_effective_referential(departement: str = "gc") -> Dict:
    """Return the active referential for the given department.

    Priority:
      1. DB-backed referential (cached, merged with fallback for GC).
      2. JSON-file generic referential from ``refs/``.
      3. Built-in ``_GC_FALLBACK_REF`` for GC, or ``_GENERIC_FALLBACK_REF``.
    """
    dept_key = departement.lower().strip()
    db_ref = _load_ref_from_db(dept_key)
    if db_ref is not None:
        return db_ref
    # GC without DB: use in-memory fallback
    if dept_key in ("gc", "genie_civil", "genie-civil"):
        return _GC_FALLBACK_REF
    # Other departments without DB → load from JSON generic ref
    logger.info(f"Chargement du référentiel générique pour le département '{dept_key}'")
    return _load_generic_ref(dept_key)


# Keep backward-compatible alias (used by older callers)
def _get_effective_gc_referential() -> Dict:
    return _get_effective_referential("gc")

# Legacy name alias: tests and external code may still import _GC_REFERENTIAL
_GC_REFERENTIAL = _GC_FALLBACK_REF


# ─────────────────────────────────────────────────────────────────────────────
# _DepartmentReferentialManager – OOP facade for multi-department referentials
# ─────────────────────────────────────────────────────────────────────────────

class _DepartmentReferentialManager:
    """High-level façade for accessing per-department referentials.

    Wraps the module-level caching functions (_load_ref_from_db,
    _get_effective_referential) and exposes a clean object-oriented API so
    callers never need to deal with raw cache globals.

    Singleton-like design: use ``_dept_ref_manager`` module-level instance.

    Supported operations:
    * ``get_referential(dept)``     → full Dict with savoirs/competences/domaines
    * ``match_savoir(text, dept)``  → List[str] of matching savoir codes
    * ``match_competence(text, dept)`` → Optional[str] best-matching competence code
    * ``get_niveau(codes, dept)``   → Optional[str] best official niveau
    * ``invalidate(dept)``         → flush cache for a single department
    * ``invalidate_all()``         → flush all department caches
    """

    # Supported ESPRIT department codes
    KNOWN_DEPARTMENTS: List[str] = ["gc", "info", "ge", "meca", "telecom"]

    def get_referential(self, department: str) -> Dict:
        """Return the effective referential for *department*.

        Loads from DB (with TTL cache) and falls back to built-in data for GC.
        Never returns None – returns _GENERIC_FALLBACK_REF as last resort.
        """
        return _get_effective_referential(department)

    def match_savoir(self, text: str, department: str) -> List[str]:
        """Match *text* against the *department* savoir referential.

        Combines semantic + keyword matching; returns deduplicated code list.
        """
        return _match_gc_savoir(text, departement=department)

    def match_competence(self, text: str, department: str) -> Optional[str]:
        """Return the best-matching competence code for *text* in *department*."""
        return _match_gc_competence(text, departement=department)

    def get_niveau(self, savoir_codes: List[str], department: str) -> Optional[str]:
        """Return the official niveau string for the highest-priority matching code."""
        return _gc_ref_niveau(savoir_codes, departement=department)

    def suggest_teachers(self, savoir_codes: List[str]) -> List[str]:
        """Return enseignant IDs mapped to any of *savoir_codes* in the DB."""
        return _suggest_gc_enseignants(savoir_codes)

    def detect_type(self, text: str, department: str) -> str:
        """Classify *text* as PRATIQUE or THEORIQUE using dept-aware patterns."""
        return _detect_type(text, departement=department)

    def invalidate(self, department: str) -> None:
        """Flush the cached referential for a single department."""
        dept_key = department.lower().strip()
        _REF_DB_CACHE.pop(dept_key)
        logger.info("Referential cache invalidated for [%s]", dept_key)

    def invalidate_all(self) -> None:
        """Flush ALL department referential caches."""
        global _SEMANTIC_CORPUS_BUILT
        _REF_DB_CACHE.clear()
        _SEMANTIC_CORPUS_BUILT = False
        logger.info("All referential caches invalidated")

    def list_departments(self) -> List[str]:
        """Return list of known ESPRIT department codes."""
        return list(self.KNOWN_DEPARTMENTS)

    def stats(self, department: str) -> Dict[str, int]:
        """Return referential size stats for *department*."""
        ref = self.get_referential(department)
        return {
            "savoirs":     len(ref.get("savoirs", {})),
            "competences": len(ref.get("competences", {})),
            "domaines":    len(ref.get("domaines", {})),
            "niveaux":     len(ref.get("niveaux", {})),
        }


# Module-level singleton: use this instance everywhere
_dept_ref_manager = _DepartmentReferentialManager()


# ─────────────────────────────────────────────────────────────────────────────
# Point 2 – Semantic GC matching (sentence-transformers, optional)
# ─────────────────────────────────────────────────────────────────────────────

_SEMANTIC_MODEL = None
_SEMANTIC_CORPUS: List[Tuple[str, Any]] = []   # (code, numpy_embedding)
_SEMANTIC_CORPUS_BUILT: bool = False
_SEMANTIC_CORPUS_DEPT: str = ""  # tracks which department the corpus was built for
# Lock for thread-safe corpus building (concurrent analysis requests)
_SEMANTIC_CORPUS_LOCK = _threading.Lock()


def _get_semantic_model():
    """Lazy-load the sentence-transformers model (downloaded once, then cached)."""
    global _SEMANTIC_MODEL
    if _SEMANTIC_MODEL is None and _SEMANTIC_OK:
        try:
            _SEMANTIC_MODEL = _SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("Semantic model loaded: all-MiniLM-L6-v2")
        except Exception as exc:
            logger.warning(f"Cannot load semantic model: {exc}")
    return _SEMANTIC_MODEL


_SEMANTIC_CACHE_DIR = Path(__file__).parent / "_semantic_cache"


def _build_semantic_corpus(departement: str = "gc") -> None:
    """Pre-compute embeddings for every savoir description for the given department.

    The corpus is invalidated and rebuilt automatically when the department changes
    or when the DB referential cache is refreshed.
    Thread-safe: protected by _SEMANTIC_CORPUS_LOCK to avoid concurrent rebuilds.

    Embeddings are persisted to disk (``_semantic_cache/<dept>.npy``) so that
    subsequent startups skip the expensive encode step.
    """
    global _SEMANTIC_CORPUS, _SEMANTIC_CORPUS_BUILT, _SEMANTIC_CORPUS_DEPT
    with _SEMANTIC_CORPUS_LOCK:
        # Double-checked: another thread may have built the corpus while we waited
        dept_key = departement.lower().strip()
        if _SEMANTIC_CORPUS_BUILT and _SEMANTIC_CORPUS_DEPT == dept_key:
            return
        model = _get_semantic_model()
        if not model:
            return
        ref = _get_effective_referential(departement)
        codes = list(ref["savoirs"].keys())
        descriptions = [" ".join(kws) for kws in ref["savoirs"].values()]
        if not codes:
            return

        # ── Try loading persisted embeddings from disk ────────────────────
        cache_file = _SEMANTIC_CACHE_DIR / f"{dept_key}.npy"
        if _SEMANTIC_OK and cache_file.is_file():
            try:
                cached_embeddings = _np.load(str(cache_file))
                if cached_embeddings.shape[0] == len(codes):
                    _SEMANTIC_CORPUS = list(zip(codes, cached_embeddings))
                    _SEMANTIC_CORPUS_BUILT = True
                    _SEMANTIC_CORPUS_DEPT = dept_key
                    logger.info(f"Embeddings sémantiques chargés depuis le disque pour [{dept_key}] "
                                f"({len(codes)} vecteurs)")
                    return
                else:
                    logger.info(f"Cache embeddings size mismatch for [{dept_key}], rebuilding")
            except Exception as exc:
                logger.warning(f"Erreur lecture cache embeddings [{dept_key}]: {exc}")

        # ── Build from scratch and persist ────────────────────────────────
        try:
            embeddings = model.encode(descriptions, convert_to_numpy=True)
            _SEMANTIC_CORPUS = list(zip(codes, embeddings))
            _SEMANTIC_CORPUS_BUILT = True
            _SEMANTIC_CORPUS_DEPT = dept_key
            logger.info(f"Semantic corpus built [{departement}]: {len(_SEMANTIC_CORPUS)} savoir embeddings")
            # Persist to disk
            try:
                _SEMANTIC_CACHE_DIR.mkdir(parents=True, exist_ok=True)
                _np.save(str(cache_file), embeddings)
                logger.info(f"Embeddings sémantiques persistés sur disque ({len(embeddings)} vecteurs) [{dept_key}]")
            except Exception as save_exc:
                logger.warning(f"Impossible de persister le corpus sémantique [{dept_key}]: {save_exc}")
        except Exception as exc:
            logger.warning(f"Cannot build semantic corpus: {exc}")


def _match_gc_savoir_semantic(text: str, threshold: float = 0.35, top_k: int = 5,
                               departement: str = "gc") -> List[str]:
    """Semantic referential matching via cosine similarity.

    Returns up to `top_k` savoir codes whose description embedding is ≥
    `threshold` similar to the query embedding.  Returns [] when
    sentence-transformers is not installed.
    """
    if not _SEMANTIC_OK:
        return []
    # Rebuild corpus if department changed or not yet built (lock inside _build)
    if not _SEMANTIC_CORPUS_BUILT or _SEMANTIC_CORPUS_DEPT != departement.lower().strip():
        _build_semantic_corpus(departement)
    if not _SEMANTIC_CORPUS:
        return []
    model = _get_semantic_model()
    if not model:
        return []
    try:
        q_emb = model.encode([text], convert_to_numpy=True)[0]
        scores: List[Tuple[str, float]] = []
        for code, emb in _SEMANTIC_CORPUS:
            norm = float(_np.linalg.norm(q_emb) * _np.linalg.norm(emb))
            sim = float(_np.dot(q_emb, emb)) / (norm + 1e-8)
            if sim >= threshold:
                scores.append((code, sim))
        scores.sort(key=lambda x: x[1], reverse=True)
        return [c for c, _ in scores[:top_k]]
    except Exception as exc:
        logger.warning(f"Semantic matching error: {exc}")
        return []


def _detect_departement(filenames: List[str], contents: List[bytes]) -> str:
    """Heuristically infer the ESPRIT department code from filenames and file text.

    Scans the first 4 KB of each file (UTF-8 best-effort) together with the
    filenames and returns one of: ``'gc'``, ``'info'``, ``'ge'``, ``'meca'``,
    ``'telecom'``.  Falls back to ``'gc'`` when no signal is found.
    """
    combined = " ".join(filenames).lower()
    for data in contents:
        try:
            combined += " " + data[:4096].decode("utf-8", errors="ignore").lower()
        except Exception:
            pass

    _DEPT_SIGNALS = [
        ("gc",     ["genie civil", "beton", "fondation", "structure portante",
                    "hydraulique", "topographie", "geotechnique", "ouvrage"]),
        ("info",   ["informatique", "algorithmique", "programmation", "logiciel",
                    "base de donnee", "reseau informatique", "developpement web"]),
        ("ge",     ["genie electrique", "electronique", "electrotechnique",
                    "automatique", "energie electrique"]),
        ("meca",   ["genie mecanique", "mecanique", "thermodynamique",
                    "fabrication", "usinage"]),
        ("telecom", ["telecom", "telecommunication", "signal numerique",
                     "radiocommunication"]),
    ]
    best_dept, best_score = "gc", 0
    for dept_code, keywords in _DEPT_SIGNALS:
        score = sum(1 for kw in keywords if kw in combined)
        if score > best_score:
            best_score = score
            best_dept = dept_code
    logger.info(f"Auto-detected department: '{best_dept}' (score={best_score})")
    return best_dept


@rice_router.post("/analyze", response_model=RiceAnalysisResult,
                  summary="Analyser les fiches UE/modules et générer le référentiel RICE")
async def rice_analyze(
    files: List[UploadFile] = File(..., description="Fiches UE et modules (PDF/DOCX)"),
    enseignants: str = Form(
        default="[]",
        description='JSON array: [{id, nom, prenom, modules:[...]}]',
    ),
    departement: str = Form(
        default="auto",
        description=(
            "Code du département ESPRIT : 'auto' (détection automatique depuis la fiche), "
            "'gc' (Génie Civil, référentiel intégré), "
            "'info' (Informatique), 'ge' (Génie Électrique), "
            "'meca' (Génie Mécanique), 'telecom' (Télécommunications), "
            "ou tout autre code défini dans la table ref_savoirs."
        ),
    ),
    _user: Optional[Dict] = Depends(_get_current_user),
):
    if not files:
        raise HTTPException(400, "Au moins un fichier est requis")

    try:
        ens_list = [EnseignantInfo(**e) for e in _json.loads(enseignants)]
    except Exception as exc:
        raise HTTPException(400, f"Format enseignants invalide: {exc}") from exc

    # Sanitize filenames to prevent path-traversal attacks
    filenames = [
        os.path.basename((f.filename or "").replace("..", "")) or f"file_{i}"
        for i, f in enumerate(files)
    ]
    contents  = [await f.read() for f in files]

    # Department resolution: "auto" → heuristic detection from filenames & content
    dept_key = departement.lower().strip()
    if dept_key == "auto":
        dept_key = _detect_departement(filenames, contents)
        departement = dept_key

    # Run CPU-bound analysis in a thread pool to avoid blocking the event loop
    result = await run_in_threadpool(analyze_files, filenames, contents, ens_list, departement)
    return result


@rice_router.get("/referential",
                 summary="Obtenir le référentiel du département pour matching")
@rice_router.get("/gc-referential",   # backward-compat alias
                 summary="Alias déprécié – utiliser /referential", include_in_schema=False)
def get_referential(
    departement: str = "gc",
):
    """Return the department referential (+ dynamic enseignant affectations).

    Use ``?departement=gc`` for Génie Civil (default),
    or ``?departement=info``, ``?departement=ge``, ``?departement=telecom``, etc.
    for other ESPRIT departments.
    """
    return {
        **_get_effective_referential(departement),
        "departement": departement,
        "enseignant_affectations": _fetch_enseignant_affectations(),
    }


@rice_router.post("/refresh-cache",
                  summary="Forcer le rafraîchissement du cache du référentiel et des affectations")
@rice_router.post("/gc-refresh-cache",  # backward-compat alias
                  summary="Alias déprécié – utiliser /refresh-cache", include_in_schema=False)
def refresh_cache(_user: Optional[Dict] = Depends(_get_current_user)):
    """Invalidate the affectations cache AND all department referential DB caches
    so the next call reads fresh data from DB."""
    global _SEMANTIC_CORPUS_BUILT
    _AFFECTATIONS_CACHE.clear()
    _REF_DB_CACHE.clear()
    _SEMANTIC_CORPUS_BUILT = False  # force re-encoding if ref changed
    fresh = _fetch_enseignant_affectations()
    return {"status": "ok", "enseignants_count": len(fresh)}


@rice_router.post("/match",
                  summary="Matcher un texte libre contre le référentiel d'un département")
@rice_router.post("/gc-match",  # backward-compat alias
                  summary="Alias déprécié – utiliser /match", include_in_schema=False)
async def match_text(
    text: str = Form(..., description="Texte à matcher (objectif, contenu de fiche…)"),
    departement: str = Form(default="gc", description="Code département (gc, info, ge, telecom, meca, …)"),
    _user: Optional[Dict] = Depends(_get_current_user),
):
    """
    Match free text against the department referential.
    Returns matched savoirs, competence, and suggested enseignants.
    Works for all ESPRIT departments: gc, info, ge, telecom, meca, …
    """
    savoir_codes = _match_gc_savoir(text, departement=departement)
    competence = _match_gc_competence(text, departement=departement)
    suggested_ens = _suggest_gc_enseignants(savoir_codes[:5])
    return {
        "departement": departement,
        "matched_savoirs": savoir_codes[:10],
        "matched_competence": competence,
        "suggested_enseignants": suggested_ens,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Point 6 – /validate  (human review → persist to DB)
# ─────────────────────────────────────────────────────────────────────────────

class ValidateRequest(BaseModel):
    """Body sent by the frontend after drag-&-drop review of the analysis result."""
    propositions: List[DomaineProposition]
    # When True, existing rows with the same id are deleted before re-insert
    overwrite: bool = False


class ValidateSummary(BaseModel):
    status: str
    upserted_domaines: int = 0
    upserted_competences: int = 0
    upserted_sous_competences: int = 0
    inserted_savoirs: int
    updated_savoirs: int
    inserted_enseignant_links: int
    errors: List[str] = []


@rice_router.post(
    "/validate",
    response_model=ValidateSummary,
    summary="Valider et persister le référentiel RICE validé par l'humain en BDD",
)
async def rice_validate(request: ValidateRequest, _user: Optional[Dict] = Depends(_get_current_user)):
    """
    Persist the human-validated competence tree to PostgreSQL.

    Traverses the full Domaine → Compétence → SousCompétence → Savoir hierarchy
    and upserts every level in order so that the foreign-key chain is always intact.

    Actions performed at each level:
    * **domaines**: ``INSERT … ON CONFLICT (code) DO UPDATE``
    * **competences**: idem, with ``domaine_id`` FK resolved by code.
    * **sous_competences**: idem, with ``competence_id`` FK resolved by code.
    * **savoirs**: idem by UUID ``id``, with ``sous_competence_id`` FK resolved by code.
    * **enseignant_competences**: ``INSERT … ON CONFLICT DO NOTHING`` for each
      suggested teacher.  If ``overwrite=True`` existing links for the savoir
      are deleted first so that the reviewer's corrections win.

    Returns a ``ValidateSummary`` with per-level counts and non-fatal errors.
    """
    if not request.propositions:
        raise HTTPException(400, "propositions list is empty")

    upserted_domaines = 0
    upserted_competences = 0
    upserted_sous_competences = 0
    inserted_savoirs = 0
    updated_savoirs = 0
    inserted_links = 0
    errors: List[str] = []

    try:
        conn = _get_db_connection()
        cur = conn.cursor()
    except Exception as exc:
        raise HTTPException(503, f"Database unreachable: {exc}") from exc

    try:
        for domaine in request.propositions:
            # ── Upsert domaine ────────────────────────────────────────────────
            try:
                cur.execute("""
                    INSERT INTO domaines (code, nom, description, actif)
                    VALUES (%s, %s, %s, true)
                    ON CONFLICT (code) DO UPDATE SET
                        nom         = EXCLUDED.nom,
                        description = EXCLUDED.description
                """, (
                    domaine.code[:50],
                    domaine.nom[:255],
                    (domaine.description or "")[:500],
                ))
                upserted_domaines += 1
            except Exception as dom_err:
                conn.rollback()
                errors.append(f"domaine {domaine.code}: {dom_err}")
                continue

            for competence in domaine.competences:
                # ── Upsert competence ─────────────────────────────────────────
                try:
                    cur.execute("""
                        INSERT INTO competences
                            (code, nom, description, ordre, domaine_id)
                        VALUES (
                            %s, %s, %s, %s,
                            (SELECT id FROM domaines WHERE code = %s)
                        )
                        ON CONFLICT (code) DO UPDATE SET
                            nom        = EXCLUDED.nom,
                            description = EXCLUDED.description,
                            ordre      = EXCLUDED.ordre,
                            domaine_id = EXCLUDED.domaine_id
                    """, (
                        competence.code[:50],
                        competence.nom[:255],
                        (competence.description or "")[:500],
                        competence.ordre,
                        domaine.code[:50],
                    ))
                    upserted_competences += 1
                except Exception as comp_err:
                    conn.rollback()
                    errors.append(f"competence {competence.code}: {comp_err}")
                    continue

                for sc in competence.sousCompetences:
                    # ── Upsert sous_competence ────────────────────────────────
                    try:
                        cur.execute("""
                            INSERT INTO sous_competences
                                (code, nom, description, competence_id)
                            VALUES (
                                %s, %s, %s,
                                (SELECT id FROM competences WHERE code = %s)
                            )
                            ON CONFLICT (code) DO UPDATE SET
                                nom          = EXCLUDED.nom,
                                description  = EXCLUDED.description,
                                competence_id = EXCLUDED.competence_id
                        """, (
                            sc.code[:50],
                            sc.nom[:255],
                            (sc.description or "")[:500],
                            competence.code[:50],
                        ))
                        upserted_sous_competences += 1
                    except Exception as sc_err:
                        conn.rollback()
                        errors.append(f"sous_competence {sc.code}: {sc_err}")
                        continue

                    for savoir in sc.savoirs:
                        sav_id = savoir.tmpId
                        # ── Upsert savoir (with sous_competence FK) ───────────
                        try:
                            cur.execute("""
                                INSERT INTO savoirs
                                    (id, code, nom, description, type, niveau,
                                     sous_competence_id)
                                VALUES (
                                    %s, %s, %s, %s, %s, %s,
                                    (SELECT id FROM sous_competences WHERE code = %s)
                                )
                                ON CONFLICT (id) DO UPDATE SET
                                    code               = EXCLUDED.code,
                                    nom                = EXCLUDED.nom,
                                    description        = EXCLUDED.description,
                                    type               = EXCLUDED.type,
                                    niveau             = EXCLUDED.niveau,
                                    sous_competence_id = EXCLUDED.sous_competence_id
                                RETURNING (xmax = 0) AS inserted
                            """, (
                                sav_id,
                                savoir.code[:50],
                                savoir.nom[:255],
                                (savoir.description or "")[:500],
                                savoir.type,
                                savoir.niveau,
                                sc.code[:50],
                            ))
                            row = cur.fetchone()
                            if row and row[0]:
                                inserted_savoirs += 1
                            else:
                                updated_savoirs += 1
                        except Exception as sav_err:
                            conn.rollback()
                            errors.append(f"savoir {savoir.code}: {sav_err}")
                            continue

                        # ── Teacher links ─────────────────────────────────────
                        if request.overwrite:
                            try:
                                cur.execute(
                                    "DELETE FROM enseignant_competences WHERE savoir_id = %s",
                                    (sav_id,),
                                )
                            except Exception:
                                conn.rollback()

                        for ens_id in savoir.enseignantsSuggeres:
                            try:
                                cur.execute("""
                                    INSERT INTO enseignant_competences
                                        (enseignant_id, savoir_id, niveau, date_acquisition)
                                    VALUES (%s, %s, %s, CURRENT_DATE)
                                    ON CONFLICT DO NOTHING
                                """, (ens_id, sav_id, savoir.niveau))
                                inserted_links += 1
                            except Exception as link_err:
                                conn.rollback()
                                errors.append(f"link {ens_id}->{savoir.code}: {link_err}")

        conn.commit()
    except Exception as exc:
        conn.rollback()
        raise HTTPException(500, f"DB error during validation: {exc}") from exc
    finally:
        try:
            cur.close()
            _put_db_connection(conn)
        except Exception:
            pass

    logger.info(
        f"Validate: domaines={upserted_domaines} competences={upserted_competences} "
        f"sous_competences={upserted_sous_competences} "
        f"inserted={inserted_savoirs} updated={updated_savoirs} "
        f"links={inserted_links} errors={len(errors)}"
    )
    return ValidateSummary(
        status="ok",
        upserted_domaines=upserted_domaines,
        upserted_competences=upserted_competences,
        upserted_sous_competences=upserted_sous_competences,
        inserted_savoirs=inserted_savoirs,
        updated_savoirs=updated_savoirs,
        inserted_enseignant_links=inserted_links,
        errors=errors[:30],
    )
